import io
import requests
import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import tempfile
import intraday_screener
import time
import streamlit.components.v1 as components
import mc_simulator as mc

from docx.shared import Inches
from docx import Document
from datetime import datetime
from docx.shared import RGBColor
from ui_themes import apply_theme, style_signal_table, theme_picker
from backtest_engine import (
    run_walk_forward_backtest,
    run_macro_strategy_backtest,
    build_portfolio_backtest,
    compute_backtest_analytics,
    MacroBacktestConfig,
)
from strategy_engine import decide_trade
from macro_fusion import build_macro_state
import plotly.graph_objects as go

from walkforward_mc_validation import (
    MCValidationConfig,
    run_walkforward_mc_validation,
    summarize_mc_validation,
    build_tactical_state_series,
)

# -----------------------------
# FRED CSV download helpers
# -----------------------------

def auto_refresh(seconds: int, enabled: bool):
    """
    Best-effort auto refresh without extra packages.
    Inserts an HTML meta-refresh tag.
    """
    if enabled and seconds and seconds > 0:
        components.html(
            f"<meta http-equiv='refresh' content='{int(seconds)}'>",
            height=0,
            width=0,
        )

def save_fig_to_tempfile(fig) -> str:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(tmp.name, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return tmp.name

def fred_csv(series_id: str) -> pd.Series:
    """
    Pull a FRED series as a pandas Series using FRED's CSV download endpoint.
    Uses /graph/fredgraph.csv?id=SERIES which typically works without an API key.
    """
    url = f"https://fred.stlouisfed.org/graph/fredgraph.csv?id={series_id}"
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    df = pd.read_csv(io.StringIO(r.text))
    # FRED CSV typically has columns: DATE, <SERIES_ID>
    date_col = df.columns[0]
    val_col = df.columns[1]
    df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
    df[val_col] = pd.to_numeric(df[val_col], errors="coerce")
    s = df.set_index(date_col)[val_col].sort_index()
    s.name = series_id
    return s

def to_monthly(s: pd.Series, method: str = "avg") -> pd.Series:
    s = s.dropna()
    if s.empty:
        return s

    rule = "ME"  # month-end (pandas recommended; avoids FutureWarning)

    if method == "eom":
        return s.resample(rule).last()

    # default "avg"
    return s.resample(rule).mean()

def splice_index(old: pd.Series, new: pd.Series, anchor_date="2006-01-31") -> pd.Series:
    """
    Splice two overlapping index series by scaling 'old' to match 'new' at anchor_date.
    """
    old_m = old.copy()
    new_m = new.copy()
    old_m.index = pd.to_datetime(old_m.index)
    new_m.index = pd.to_datetime(new_m.index)

    # Find nearest available anchor in each series
    anchor = pd.to_datetime(anchor_date)
    old_anchor = old_m.loc[:anchor].dropna().iloc[-1] if not old_m.loc[:anchor].dropna().empty else np.nan
    new_anchor = new_m.loc[:anchor].dropna().iloc[-1] if not new_m.loc[:anchor].dropna().empty else np.nan

    if not (np.isfinite(old_anchor) and np.isfinite(new_anchor) and old_anchor != 0):
        # Fallback: just concat with preference for new
        combined = pd.concat([old_m, new_m]).sort_index()
        combined = combined[~combined.index.duplicated(keep="last")]
        return combined

    scale = new_anchor / old_anchor
    old_scaled = old_m * scale
    combined = pd.concat([old_scaled, new_m]).sort_index()
    combined = combined[~combined.index.duplicated(keep="last")]
    return combined

def last_update_date(df: pd.DataFrame, col: str) -> str:
    s = df[col].dropna()
    if s.empty:
        return "—"
    return s.index.max().date().isoformat()

def latest_value_and_date(df: pd.DataFrame, col: str):
    s = df[col].dropna()
    if s.empty:
        return (np.nan, "—")
    return (float(s.iloc[-1]), s.index.max().date().isoformat())

def rsi(series: pd.Series, period: int = 14) -> pd.Series:
    """
    Wilder RSI. Assumes a regular time series (monthly in our use case).
    Returns RSI in [0,100].
    """
    s = series.astype(float)
    delta = s.diff()
    gain = delta.clip(lower=0.0)
    loss = (-delta).clip(lower=0.0)

    # Wilder smoothing
    avg_gain = gain.ewm(alpha=1/period, adjust=False, min_periods=period).mean()
    avg_loss = loss.ewm(alpha=1/period, adjust=False, min_periods=period).mean()

    rs = avg_gain / avg_loss.replace(0.0, np.nan)
    out = 100.0 - (100.0 / (1.0 + rs))
    return out

def zscore(series: pd.Series, window: int = 60) -> pd.Series:
    """
    Rolling z-score using trailing window.
    For RSI regime-adaptation: default window=60 months (~5 years).
    """
    s = series.astype(float)
    mu = s.rolling(window).mean()
    sd = s.rolling(window).std(ddof=0)
    return (s - mu) / sd.replace(0.0, np.nan)

def pct_slope(series: pd.Series, lookback: int) -> pd.Series:
    s = pd.to_numeric(series, errors="coerce")
    return (s / s.shift(lookback) - 1.0) * 100.0

# =============================
# Intraday (yfinance) helpers
# =============================

@st.cache_data(ttl=65, show_spinner=False)
def yf_intraday_close(
    tickers: list[str],
    interval: str = "15m",
    lookback_days: int = 10,
    refresh_token: int = 0,
) -> pd.Series:
    """
    Best-effort intraday close series via yfinance.
    Returns a tz-naive pandas Series named GOLD_INTRA.
    CACHED (ttl ~5 minutes) to avoid repeated calls.
    """
    try:
        import yfinance as yf
    except Exception:
        return pd.Series(dtype=float)

    # yfinance supports period strings like "5d", "10d", "1mo"
    if lookback_days <= 5:
        period = "5d"
    elif lookback_days <= 10:
        period = "10d"
    elif lookback_days <= 30:
        period = "1mo"
    else:
        # safe cap
        period = "3mo"

    for t in tickers:
        try:
            df_i = yf.download(
                t,
                period=period,
                interval=interval,
                progress=False,
                auto_adjust=True,
                prepost=False,
            )
            if df_i is None or df_i.empty:
                continue

            col = "Close" if "Close" in df_i.columns else ("Adj Close" if "Adj Close" in df_i.columns else None)
            if col is None:
                continue

            s = df_i[col]
            # yfinance sometimes returns DataFrame for single ticker; flatten
            if isinstance(s, pd.DataFrame):
                s = s.iloc[:, 0]

            s = s.dropna()
            if s.empty:
                continue

            # Make tz-naive index for consistent plotting
            try:
                if getattr(s.index, "tz", None) is not None:
                    s.index = s.index.tz_convert(None)
            except Exception:
                pass

            s.name = "GOLD_INTRA"
            return s.sort_index()

        except Exception:
            continue

    return pd.Series(dtype=float)


def compute_intraday_rsi_pack(
    intraday_close: pd.Series,
    rsi_period: int = 14,
) -> dict:
    """
    Returns dict with RSI series + optional diagnostics.
    Uses your existing Wilder RSI function.
    """
    if not isinstance(intraday_close, pd.Series) or intraday_close.dropna().empty:
        return {}

    s = intraday_close.dropna().astype(float).sort_index()
    r = rsi(s, period=rsi_period).dropna()
    if r.empty:
        return {}

    pack = {
        "RSI_14_INTRA": r,
        # lightweight extras (optional but handy)
        "RSI_14_INTRA_Z_2W": zscore(r, window=14 * 10),   # ~2 weeks of 5 trading days/week * ~10 bars/day (rough)
        "RSI_14_INTRA_SLOPE_3H": r.diff(36),              # rough: 36 * 5m = 3h (approx); ok even for 15m (then ~9h)
    }
    return pack

def classify_rsi_overlay(rsi_val: float, rsi_z: float, rsi_slope_3m: float) -> str:
    """
    Friendly overlay label for the Acceleration dashboard/report.
    Uses z-scored RSI + RSI slope (momentum cooling/strengthening).
    """
    if pd.isna(rsi_val):
        return "RSI unavailable"

    # Stretchedness based on z-score (regime-adaptive)
    if not pd.isna(rsi_z):
        if rsi_z >= 1.5:
            stretch = "Stretched (high)"
        elif rsi_z >= 1.0:
            stretch = "Elevated"
        elif rsi_z <= -1.5:
            stretch = "Stretched (low)"
        elif rsi_z <= -1.0:
            stretch = "Depressed"
        else:
            stretch = "Normal"
    else:
        # fallback if z-score not available yet
        if rsi_val >= 70:
            stretch = "Overbought"
        elif rsi_val <= 30:
            stretch = "Oversold"
        else:
            stretch = "Neutral"

    # Momentum direction from slope
    if pd.isna(rsi_slope_3m):
        mom = "Momentum unknown"
    else:
        if rsi_slope_3m >= 5:
            mom = "Momentum strengthening"
        elif rsi_slope_3m <= -5:
            mom = "Momentum cooling"
        else:
            mom = "Momentum steady"

    return f"{stretch} | {mom}"

# -----------------------------
# Model construction
# -----------------------------
def build_features(monthly_method: str = "avg") -> pd.DataFrame:

    # Rates
    dgs10 = to_monthly(fred_csv("DGS10"), monthly_method)  # daily -> monthly
    tb3ms = to_monthly(fred_csv("TB3MS"), "avg")           # already monthly, avg ok

    # Inflation: CPI YoY (proxy for inflation expectations when breakevens unavailable)
    cpi = to_monthly(fred_csv("CPIAUCSL"), "avg")
    cpi_yoy = 100 * (cpi / cpi.shift(12) - 1.0)
    cpi_yoy.name = "CPI_YOY"

    real_yield_cpi = dgs10 - cpi_yoy
    real_yield_cpi.name = "REAL_YIELD_CPI"

    # Breakevens: 10Y (available mainly from early 2000s onward)
    t10yie = to_monthly(fred_csv("T10YIE"), monthly_method)
    t10yie.name = "T10YIE"

    # Composite "inflation expectations" proxy:
    # prefer breakevens when available; otherwise fall back to CPI YoY
    infl_exp = t10yie.combine_first(cpi_yoy)
    infl_exp.name = "INFL_EXP_PROXY"

    # Real yield proxy (long-history): 10y nominal - CPI YoY
    real_yield_proxy = dgs10 - infl_exp
    real_yield_proxy.name = "REAL_YIELD_PROXY"

    # Dollar index: splice long history (TWEXBMTH) with modern (TWEXBGSMTH)
    twexbmth = to_monthly(fred_csv("TWEXBMTH"), "avg")     # 1973–2019
    twexbgsmth = to_monthly(fred_csv("TWEXBGSMTH"), "avg") # 2006–present
    usd_idx = splice_index(twexbmth, twexbgsmth)
    usd_idx.name = "USD_TWEX_SPLICE"

    usd_12m_chg = 100 * (usd_idx / usd_idx.shift(12) - 1.0)
    usd_12m_chg.name = "USD_12M_CHG"

    # Curve proxy (long-history): 10Y - 3M (monthly)
    curve = dgs10 - tb3ms
    curve.name = "CURVE_10Y_3M"

    # Fiscal dominance proxy: deficit % GDP (often quarterly/annual -> ffill monthly)
    deficit_gdp = fred_csv("FYFSGDA188S")
    deficit_gdp = deficit_gdp.resample("ME").ffill()
    deficit_gdp.name = "DEFICIT_GDP"

    # --- NEW: Market real yield (TIPS) + Stress premium (HY OAS)
    # TIPS 10Y real yield (daily). Not available pre-2000s.
    dfii10 = fred_csv("DFII10")
    dfii10_20d = dfii10.rolling(20).mean()
    real_yield_tips10 = dfii10_20d.dropna().resample("ME").last()
    real_yield_tips10.name = "REAL_YIELD_TIPS10"

    # High Yield OAS (daily). History starts later than 1970s.
    hy_oas = fred_csv("BAMLH0A0HYM2")
    hy_oas_20d = hy_oas.rolling(20).mean()
    hy_oas_m = to_monthly(hy_oas_20d, monthly_method)
    hy_oas_m.name = "HY_OAS"

    # Optional gold for history charts (safe if missing)
    gold = None
    gold_source = None
    gold_err = None
    gold_daily = None

    # (A) FRED (LBMA) — removed from FRED since 2022-01-31, so likely to fail
    for sid in ["GOLDAMGBD228NLBM", "GOLDPMGBD228NLBM"]:
        try:
            gold = to_monthly(fred_csv(sid), "avg")
            gold.name = "GOLD_USD"
            gold_source = f"FRED:{sid}"
            break
        except Exception as e:
            gold_err = f"{sid}: {type(e).__name__} — {e}"

    # (B) Fallback: yfinance (works without FRED)
    if gold is None:
        try:
            import yfinance as yf

            # Prefer more reliable Yahoo symbols over XAUUSD=X
            for ticker in ["GC=F", "GLD"]:
                s = yf.download(
                    ticker,
                    start="1970-01-01",
                    progress=False,
                    auto_adjust=True
                )
                # yfinance returns a DataFrame; use Close/Adj Close depending on availability
                if not s.empty:
                    col = "Close" if "Close" in s.columns else ("Adj Close" if "Adj Close" in s.columns else None)
                    if col is None:
                        continue
                    obj = s[col]

                    # yfinance can return a DataFrame (MultiIndex columns) even for one ticker.
                    # Force daily gold into a 1D Series for RSI(14D).
                    if isinstance(obj, pd.DataFrame):
                        if obj.shape[1] == 1:
                            ser = obj.iloc[:, 0].dropna()
                        else:
                            # take first column as fallback
                            ser = obj.iloc[:, 0].dropna()
                    else:
                        ser = obj.dropna()

                    if not ser.empty:
                        # Keep DAILY for RSI(14D)
                        gold_daily = ser.copy()
                        gold_daily.name = "GOLD_USD_D"

                        # Keep MONTHLY for the rest of the dashboard
                        ser_m = ser.resample("ME").last()  # month-end
                        ser_m.name = "GOLD_USD"

                        gold = ser_m
                        gold_source = f"yfinance:{ticker}"
                        gold_err = None
                        break


            if gold is None:
                raise ValueError("yfinance returned no data for all fallback tickers (GC=F, GLD).")

        except Exception as e:
            gold_err = f"yfinance: {type(e).__name__} — {e}"

    # Save diagnostics for UI
    df = pd.concat(
        [real_yield_cpi, real_yield_proxy, infl_exp, cpi_yoy, usd_idx, usd_12m_chg, curve, deficit_gdp, real_yield_tips10, hy_oas_m],
        axis=1
    )

    if "REAL_YIELD_TIPS10" in df.columns:
        df["REAL_YIELD_TIPS10"] = df["REAL_YIELD_TIPS10"].ffill()
    if "HY_OAS" in df.columns:
        df["HY_OAS"] = df["HY_OAS"].ffill()

    if gold is not None:
        # Normalize gold to a Series (some sources may return DataFrame)
        if isinstance(gold, pd.DataFrame):
            if gold.shape[1] == 1:
                gold = gold.iloc[:, 0]
            else:
                # pick a sensible column if it exists, otherwise take the first
                for c in ["Close", "Adj Close", "GOLD_USD"]:
                    if c in gold.columns:
                        gold = gold[c]
                        break
                else:
                    gold = gold.iloc[:, 0]

        gold = gold.rename("GOLD_USD")  # Series.rename(name)

    # ✅ ALWAYS do these (even if gold later fails)
    for col in ["DEFICIT_GDP", "REAL_YIELD_TIPS10", "HY_OAS"]:
        if col in df.columns:
            df[col] = df[col].ffill()

    # add gold only if available
    if gold is not None and not gold.empty:
        df = pd.concat([df, gold], axis=1)

    # --- Market-specific indicators (QQQ trend state) ---
    try:
        import yfinance as yf

        qqq = yf.download(
            "QQQ",
            start="2000-01-01",
            progress=False,
            auto_adjust=True
        )

        qqq_close = None

        if qqq is not None and not qqq.empty:
            # yfinance may return:
            # 1) normal columns: Open/High/Low/Close/Volume
            # 2) MultiIndex columns for a single ticker
            if isinstance(qqq.columns, pd.MultiIndex):
                # try to extract the "Close" level
                if "Close" in qqq.columns.get_level_values(0):
                    qqq_close = qqq["Close"]
                    if isinstance(qqq_close, pd.DataFrame):
                        qqq_close = qqq_close.iloc[:, 0]
            else:
                if "Close" in qqq.columns:
                    qqq_close = qqq["Close"]

        if qqq_close is not None:
            qqq_close = pd.to_numeric(qqq_close, errors="coerce").dropna()
            qqq_close.index = pd.to_datetime(qqq_close.index, errors="coerce")
            qqq_close = qqq_close[qqq_close.index.notna()].sort_index()

            # Make index tz-naive safely
            try:
                if getattr(qqq_close.index, "tz", None) is not None:
                    qqq_close.index = qqq_close.index.tz_convert(None)
            except Exception:
                pass

            # Daily trend features → month-end aligned
            qqq_m = qqq_close.resample("ME").last()
            qqq_ma200 = qqq_close.rolling(200).mean().resample("ME").last()
            qqq_ma50 = qqq_close.rolling(50).mean()
            qqq_ma50_slope_20d = pct_slope(qqq_ma50, 20).resample("ME").last()

            qqq_above_ma200 = (qqq_m > qqq_ma200).astype(float)

            df["QQQ_ABOVE_MA200"] = qqq_above_ma200.reindex(df.index)
            df["QQQ_MA50_SLOPE_20D"] = qqq_ma50_slope_20d.reindex(df.index)
        else:
            df["QQQ_ABOVE_MA200"] = np.nan
            df["QQQ_MA50_SLOPE_20D"] = np.nan

    except Exception as e:
        print("QQQ feature block error:", type(e).__name__, e)
        df["QQQ_ABOVE_MA200"] = np.nan
        df["QQQ_MA50_SLOPE_20D"] = np.nan

    # --- Market breadth: % of selected universe above 200D MA ---
    try:
        breadth_tickers = [
            "AAPL", "MSFT", "NVDA", "AMZN", "META", "GOOGL", "TSLA", "AVGO",
            "COST", "AMD", "WMT", "ASML", "MU", "NFLX", "PLTR", "CSCO",
            "AMAT", "LRCX", "PEP", "INTC"
        ]

        import yfinance as yf
        breadth_raw = yf.download(
            breadth_tickers,
            start="2000-01-01",
            progress=False,
            auto_adjust=True,
            group_by="column",
            threads=True,
        )

        breadth_close = None

        if breadth_raw is not None and not breadth_raw.empty:
            if isinstance(breadth_raw.columns, pd.MultiIndex):
                if "Close" in breadth_raw.columns.get_level_values(0):
                    breadth_close = breadth_raw["Close"].copy()
            else:
                # fallback unlikely, but keep safe
                if "Close" in breadth_raw.columns:
                    breadth_close = breadth_raw[["Close"]].copy()

        if breadth_close is not None and not breadth_close.empty:
            breadth_close.index = pd.to_datetime(breadth_close.index, errors="coerce")
            breadth_close = breadth_close[breadth_close.index.notna()].sort_index()

            try:
                if getattr(breadth_close.index, "tz", None) is not None:
                    breadth_close.index = breadth_close.index.tz_convert(None)
            except Exception:
                pass

            ma200_panel = breadth_close.rolling(200).mean()
            above_ma200_panel = (breadth_close > ma200_panel).astype(float)

            breadth_pct = above_ma200_panel.mean(axis=1) * 100.0
            breadth_pct_m = breadth_pct.resample("ME").last()

            df["MARKET_BREADTH_ABOVE_MA200"] = breadth_pct_m.reindex(df.index)
        else:
            df["MARKET_BREADTH_ABOVE_MA200"] = np.nan

    except Exception as e:
        print("Breadth feature block error:", type(e).__name__, e)
        df["MARKET_BREADTH_ABOVE_MA200"] = np.nan

    print("QQQ_ABOVE_MA200 non-null:", df["QQQ_ABOVE_MA200"].notna().sum())
    print("QQQ_MA50_SLOPE_20D non-null:", df["QQQ_MA50_SLOPE_20D"].notna().sum())
    print(df[["QQQ_ABOVE_MA200", "QQQ_MA50_SLOPE_20D"]].tail())

      # --- RSI overlay (Acceleration diagnostic; not scored)

    if "GOLD_USD" in df.columns:
        df["RSI_14"] = rsi(df["GOLD_USD"], period=14)
        df["RSI_Z_60"] = zscore(df["RSI_14"], window=60)
        df["RSI_SLOPE_3M"] = df["RSI_14"] - df["RSI_14"].shift(3)

    # --- Monthly RSI "as-of-today" (uses today's daily close as provisional current-month point)
    # This keeps the 14-month interpretation but updates it with the latest daily price.
    if "GOLD_USD" in df.columns and isinstance(gold_daily, pd.Series) and not gold_daily.dropna().empty:
        monthly = df["GOLD_USD"].dropna()

        gd = gold_daily.dropna()
        today_close = float(gd.iloc[-1])
        today_dt = gd.index[-1]

        # Monthly-like series: month-ends + latest daily point (provisional current month)
        monthly_like = monthly.copy()
        monthly_like.loc[today_dt] = today_close
        monthly_like = monthly_like.sort_index()

        rsi_asof_raw = rsi(monthly_like, period=14)

        # Store as a column aligned to the monthly panel so KPI can read latest monthly row
        df["RSI_14_ASOF"] = rsi_asof_raw.reindex(df.index, method="ffill")
        df.attrs["RSI_14_ASOF_RAW"] = rsi_asof_raw

    # --- Daily RSI overlay (14D) (only if daily gold exists; not scored)

    if isinstance(gold_daily, pd.Series) and not gold_daily.dropna().empty:
        rsi_d_raw = rsi(gold_daily, period=14)
        df["RSI_14D"] = rsi_d_raw.reindex(df.index, method="ffill")
      # 1Y z-score on daily RSI (approx 252 trading days)
        df["RSI_14D_Z_1Y"] = zscore(rsi_d_raw, window=252).reindex(df.index, method="ffill")
      # ~1M slope on daily RSI (20 trading days)
        df["RSI_14D_SLOPE_1M"] = rsi_d_raw.diff(20).reindex(df.index, method="ffill")

    df.attrs["gold_error"] = gold_err

    # Store a clean datetime-indexed copy (prevents float/ordinal index issues later)
    if isinstance(gold_daily, pd.Series) and (not gold_daily.empty):
        gd = gold_daily.copy()
        gd.index = pd.to_datetime(gd.index, errors="coerce")
        gd = gd[gd.index.notna()].sort_index()
        df.attrs["gold_daily"] = gd
    else:
        df.attrs["gold_daily"] = gold_daily


    # --- Daily RSI overlay (true daily series stored in attrs; not scored)
    if isinstance(gold_daily, pd.Series) and not gold_daily.dropna().empty:
        rsi_d_raw = rsi(gold_daily, period=14).dropna()
        df.attrs["RSI_14D_DAILY"] = rsi_d_raw
        df.attrs["RSI_14D_Z_1Y_DAILY"] = zscore(rsi_d_raw, window=252)
        df.attrs["RSI_14D_SLOPE_1M_DAILY"] = rsi_d_raw.diff(20)

    # Only require non-fiscal series to align, do not stop input on the date of the smaller data set
    core_no_fiscal = ["REAL_YIELD_PROXY", "INFL_EXP_PROXY", "USD_TWEX_SPLICE", "CURVE_10Y_3M"]
    df = df.dropna(subset=core_no_fiscal, how="any")
    return df

def fwd_return_months(gold: pd.Series, months: int) -> pd.Series:
    """
    Forward % return over `months` months using month-end series.
    Return aligned to the start month (t -> t+months).
    """
    g = gold.dropna()
    return 100.0 * (g.shift(-months) / g - 1.0)

def crisis_conditioned_gold_stats(
    df: pd.DataFrame,
    win_start: str,
    win_end: str,
    thresholds: dict,
    weights: dict,
    dirs: dict,
    horizon_months: int = 6,
    band: float = 0.15,
) -> dict:
    """
    Compute gold forward-return stats inside the crisis window for months
    whose SIGNAL is within +/- band of today's SIGNAL (using crisis thresholds).
    """
    if "GOLD_USD" not in df.columns:
        return {"ok": False, "reason": "GOLD_USD not available"}

    # compute SIGNAL under crisis thresholds on the whole df
    scored = compute_signal(df, thresholds, weights, dirs)

    today_sig = float(scored["SIGNAL"].dropna().iloc[-1])
    w = scored.loc[pd.to_datetime(win_start):pd.to_datetime(win_end)].copy()
    if w.empty:
        return {"ok": False, "reason": "No data in crisis window"}

    # forward returns within the crisis window
    w["GOLD_FWD_RET"] = fwd_return_months(w["GOLD_USD"], horizon_months)

    # restrict to macro-similar months (signal proximity)
    sel = w[(w["SIGNAL"] >= today_sig - band) & (w["SIGNAL"] <= today_sig + band)].copy()
    sel = sel.dropna(subset=["GOLD_FWD_RET"])

    if sel.empty or sel["GOLD_FWD_RET"].dropna().shape[0] < 6:
        return {
            "ok": False,
            "reason": "Too few similar months with forward gold returns",
            "today_signal": today_sig,
            "n_samples": int(sel.shape[0]),
        }

    s = sel["GOLD_FWD_RET"].dropna()
    return {
        "ok": True,
        "today_signal": today_sig,
        "n_samples": int(s.shape[0]),
        "horizon_m": int(horizon_months),
        "mean_ret": float(s.mean()),
        "median_ret": float(s.median()),
        "p_pos": float((s > 0).mean()),
        "q25": float(s.quantile(0.25)),
        "q75": float(s.quantile(0.75)),
    }

def compute_thresholds_from_window(df, start, end, keys, dirs, min_points: int = 24) -> dict:
    start_dt = pd.to_datetime(start)
    end_dt = pd.to_datetime(end)
    w = df.loc[start_dt:end_dt].copy()
    if w.empty:
        return {}

    def q(series, p):
        s = series.dropna()
        return float(s.quantile(p))

    thr = {}
    for col in keys:
        if col not in w.columns:
            continue
        s = w[col].dropna()
        if len(s) < min_points:
            # Skip indicators that don't exist (or are too sparse) in this crisis window
            continue

        higher_is_bullish = bool(dirs.get(col, False))
        if higher_is_bullish:
            thr[col] = {"bull": q(s, 0.67), "bear": q(s, 0.33)}
        else:
            thr[col] = {"bull": q(s, 0.33), "bear": q(s, 0.67)}
    return thr


def structural_thresholds() -> dict:
    return {
        "REAL_YIELD_CPI": {"bull": 1.0, "bear": 2.0},     # bull if ≤ 1.0, bear if ≥ 2.0
        "CPI_YOY":        {"bull": 3.5, "bear": 2.0},     # bull if ≥ 3.5, bear if ≤ 2.0
        "USD_12M_CHG":    {"bull": -3.0, "bear": 3.0},    # bull if ≤ -3%, bear if ≥ +3%
        "CURVE_10Y_3M":   {"bull": 0.0, "bear": 1.5},     # bull if ≤ 0 (inverted), bear if ≥ 1.5
        "DEFICIT_GDP":    {"bull": -6.0, "bear": -3.0},   # bull if ≤ -6, bear if ≥ -3
        "REAL_YIELD_TIPS10": {"bull": 0.75, "bear": 1.75},  # bull if <= 0.75, bear if >= 1.75
        "HY_OAS": {"bull": 4.5, "bear": 3.2},  # bull if >= 4.5, bear if <= 3.2 (higher bullish)
    }

def market_structural_thresholds() -> dict:
    return {
        "REAL_YIELD_CPI": {"bull": 1.5, "bear": 2.5},
        "USD_12M_CHG": {"bull": -2.0, "bear": 5.0},
        "CURVE_10Y_3M": {"bull": 1.0, "bear": -0.25},
        "HY_OAS": {"bull": 3.5, "bear": 5.0},
        "QQQ_ABOVE_MA200": {"bull": 1.0, "bear": 0.0},
        "QQQ_MA50_SLOPE_20D": {"bull": 0.0, "bear": -2.0},
        "MARKET_BREADTH_ABOVE_MA200": {"bull": 60.0, "bear": 40.0},
    }

def compute_accel_features(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()

    # Gold 6M return (%)
    if "GOLD_USD" in out.columns:
        out["GOLD_6M_RET"] = 100.0 * out["GOLD_USD"].pct_change(6)
    else:
        out["GOLD_6M_RET"] = np.nan

    # Real yield 3M change (pp): prefer TIPS, fallback to CPI real yield
    if "REAL_YIELD_TIPS10" in out.columns and out["REAL_YIELD_TIPS10"].dropna().size > 0:
        ry = out["REAL_YIELD_TIPS10"]
    else:
        ry = out["REAL_YIELD_CPI"] if "REAL_YIELD_CPI" in out.columns else pd.Series(index=out.index, dtype=float)

    out["REALYIELD_3M_CHG"] = ry.diff(3)

    # Stress 3M change (pp)
    out["STRESS_3M_CHG"] = out["HY_OAS"].diff(3) if "HY_OAS" in out.columns else np.nan

    # USD 3M % change
    out["USD_3M_CHG"] = 100.0 * out["USD_TWEX_SPLICE"].pct_change(3) if "USD_TWEX_SPLICE" in out.columns else np.nan

    return out


def accel_thresholds_fixed() -> dict:
    return {
        "GOLD_6M_RET":      {"bull":  8.0,  "bear": -8.0},   # % return
        "REALYIELD_3M_CHG": {"bull": -0.30, "bear":  0.30},  # pp (lower is bullish)
        "STRESS_3M_CHG":    {"bull":  0.60, "bear": -0.60},  # pp (higher is bullish)
        "USD_3M_CHG":       {"bull": -3.0,  "bear":  3.0},   # % (lower USD is bullish)
    }


def compute_accel_thresholds_quantiles(df: pd.DataFrame, keys: list[str], dirs: dict, start: str = "2000-01-01") -> dict:
    w = df.loc[pd.to_datetime(start):].copy()
    thr = {}
    for k in keys:
        if k not in w.columns:
            continue
        s = w[k].dropna()
        if len(s) < 60:
            continue

        higher_is_bullish = bool(dirs.get(k, False))
        if higher_is_bullish:
            thr[k] = {"bull": float(s.quantile(0.67)), "bear": float(s.quantile(0.33))}
        else:
            thr[k] = {"bull": float(s.quantile(0.33)), "bear": float(s.quantile(0.67))}
    return thr

def state_score(value: float, bull: float, bear: float, higher_is_bullish: bool) -> int:
    if higher_is_bullish:
        if value >= bull:
            return +1
        if value <= bear:
            return -1
        return 0
    else:
        if value <= bull:
            return +1
        if value >= bear:
            return -1
        return 0

def compute_signal(df, thresholds, weights, dirs) -> pd.DataFrame:
    out = df.copy()

    for col, higher_bull in dirs.items():
        bull = thresholds[col]["bull"]
        bear = thresholds[col]["bear"]
        out[col + "_STATE"] = out[col].apply(lambda v: state_score(v, bull, bear, higher_bull))

    # Weighted contribution
    contrib_cols = []
    for col in dirs.keys():
        ccol = col + "_CONTRIB"
        out[ccol] = weights[col] * out[col + "_STATE"]
        contrib_cols.append(ccol)

    out["SIGNAL"] = out[contrib_cols].sum(axis=1)

    # Regime labels (editable if you want later)
    def label(sig):
        if sig >= 0.60:
            return "Structural Bull"
        if sig >= 0.20:
            return "Positive"
        if sig > -0.20:
            return "Neutral"
        if sig > -0.60:
            return "Vulnerable"
        return "Structural Headwind"

    out["REGIME"] = out["SIGNAL"].apply(label)
    return out

def add_indicator_threshold_lines(ax, ind_key: str, thresholds: dict, mode: str, color_bull="green", color_bear="red"):
    """
    Draw bull/bear threshold reference lines for an indicator, if available.
    thresholds[ind_key] expected format: {"bull": x, "bear": y} or (bull, bear).
    """
    if thresholds is None or ind_key not in thresholds:
        return

    thr = thresholds[ind_key]
    if isinstance(thr, dict):
        bull = thr.get("bull", None)
        bear = thr.get("bear", None)
    else:
        # tuple/list fallback
        bull, bear = thr[0], thr[1]

    # Only draw if numeric
    if bull is not None and pd.notna(bull):
        ax.axhline(float(bull), linestyle="--", linewidth=1.0, color=color_bull, alpha=0.85, label="Bull threshold")
    if bear is not None and pd.notna(bear):
        ax.axhline(float(bear), linestyle="--", linewidth=1.0, color=color_bear, alpha=0.85, label="Bear threshold")

def build_indicator_figure(res_plot: pd.DataFrame,
                           ind_key: str,
                           thresholds: dict,
                           mode: str,
                           labels_dict: dict):
    """
    Build a matplotlib figure for one indicator history + optional bull/bear thresholds.
    Uses res_plot (already sliced to the report window).
    """
    if ind_key not in res_plot.columns:
        return None

    s = res_plot[ind_key].dropna()
    if s.empty:
        return None

    fig, ax = plt.subplots(figsize=(8, 2.6))
    ax.plot(s.index, s.values, linewidth=1.5, label=labels_dict.get(ind_key, ind_key))

    # Threshold lines (if exist)
    try:
        add_indicator_threshold_lines(ax, ind_key, thresholds, mode)
    except Exception:
        pass

    ax.set_xlabel("Date")
    ax.set_ylabel(labels_dict.get(ind_key, ind_key))
    ax.grid(True, alpha=0.35)
    ax.legend(loc="best", fontsize=8)
    plt.tight_layout()
    return fig

def build_signal_figure(res_plot):
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(res_plot.index, res_plot["SIGNAL"], label="Gold Signal", color=theme.signal_gold)
    ax.axhline(trig_hi, linestyle="--", color="gray", label="Bull trigger")
    ax.axhline(trig_lo, linestyle="--", color="gray", label="Bear trigger")
    ax.axhline(0.0, linewidth=0.8, color="black")

    ax.set_xlabel("Date")
    ax.set_ylabel("Signal")
    ax.legend()
    ax.grid(True, alpha=0.4)
    return fig

def build_gold_signal_figure(res_plot):
    gold_col = None
    for c in ["GOLD_USD", "GC=F", "GLD"]:
        if c in res_plot.columns:
            gold_col = c
            break

    if not gold_col:
        return None

    plot_df = res_plot[[gold_col, "SIGNAL"]].dropna()
    if plot_df.empty:
        return None

    fig, ax1 = plt.subplots(figsize=(8, 4))
    l1 = ax1.plot(plot_df.index, plot_df[gold_col], color="gold", label="Gold (USD)")
    ax1.set_ylabel("Gold price (USD)")
    ax1.set_xlabel("Date")

    ax2 = ax1.twinx()
    l2 = ax2.plot(plot_df.index, plot_df["SIGNAL"], color="tab:blue", label="Signal")
    ax2.set_ylabel("Signal")

    lines = l1 + l2
    labels_ = [l.get_label() for l in lines]
    ax1.legend(lines, labels_)
    ax1.grid(True, alpha=0.4)

    return fig

def build_contrib_figure(res_plot):
    contrib_cols = [c for c in res_plot.columns if c.endswith("_CONTRIB")]
    contrib = res_plot[contrib_cols].dropna(how="all")

    if contrib.empty:
        return None

    fig, ax = plt.subplots(figsize=(8, 4))
    stacks = ax.stackplot(contrib.index, contrib.T.values)

    base_keys = [c.replace("_CONTRIB", "") for c in contrib_cols]
    leg_labels = [labels.get(k, k) for k in base_keys]

    ax.legend(stacks, leg_labels, fontsize=8)
    ax.axhline(0.0, color="black")
    ax.set_xlabel("Date")
    ax.set_ylabel("Contribution")
    ax.grid(True, alpha=0.4)

    return fig

def build_word_report_compare(
    left_title: str, right_title: str,
    left_pack: tuple, right_pack: tuple,
    trigger_info: dict,
    labels_left: dict, labels_right: dict,
    legend_map: dict
) -> bytes:
    """
    left_pack/right_pack = (res, core_keys, thresholds, dirs, weights)
    """
    (resL, coreL, thrL, dirsL, wL) = left_pack
    (resR, coreR, thrR, dirsR, wR) = right_pack

    doc = Document()
    doc.add_heading("Gold Macro Cockpit — Comparison Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    def add_mode_section(title: str, res_panel: pd.DataFrame, core_keys: list, thresholds: dict, weights: dict, labels_dict: dict):
        latest_row = res_panel.iloc[-1]
        prev_row = res_panel.iloc[-2] if len(res_panel) > 1 else latest_row

        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Latest data point: {latest_row.name.date()}")
        doc.add_paragraph(f"Signal: {float(latest_row['SIGNAL']):.2f} (Δ {float(latest_row['SIGNAL'] - prev_row['SIGNAL']):+.2f})")
        doc.add_paragraph(f"Regime: {str(latest_row.get('REGIME','—'))}")

        # Weights table
        doc.add_heading("Weights", level=3)
        t = doc.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "Indicator"
        t.rows[0].cells[1].text = "Weight"
        for k in core_keys:
            row = t.add_row().cells
            row[0].text = legend_map.get(k, labels_dict.get(k, k))
            row[1].text = f"{weights.get(k, 0.0):.2f}"

        # States table
        doc.add_heading("Indicator States (latest)", level=3)
        table = doc.add_table(rows=1, cols=5)
        h = table.rows[0].cells
        h[0].text = "Indicator"
        h[1].text = "Latest"
        h[2].text = "State"
        h[3].text = "Contribution"
        h[4].text = "Thresholds (bull/bear)"
        for col in core_keys:
            row = table.add_row().cells
            row[0].text = labels_dict.get(col, col)
            row[1].text = f"{float(latest_row[col]):.2f}" if pd.notna(latest_row[col]) else "—"
            row[2].text = str(int(latest_row.get(col + "_STATE", 0)))
            row[3].text = f"{float(latest_row.get(col + '_CONTRIB', 0.0)):+.2f}"
            row[4].text = f"{thresholds[col]['bull']:.2f} / {thresholds[col]['bear']:.2f}"

        # Indicator charts (below indicator states)
        doc.add_heading("Indicator Charts (with thresholds)", level=3)

        for k in core_keys:
            figk = build_indicator_figure(res_panel, k, thresholds, title, labels_dict)
            if figk is None:
                continue
            imgk = save_fig_to_tempfile(figk)
            doc.add_paragraph(labels_dict.get(k, k))
            doc.add_picture(imgk, width=Inches(6))

        # Charts
        doc.add_heading("Charts", level=3)

        fig1 = build_signal_figure(res_panel)
        img1 = save_fig_to_tempfile(fig1)
        doc.add_paragraph("Gold Signal over time")
        doc.add_picture(img1, width=Inches(6))

        fig2 = build_gold_signal_figure(res_panel)
        if fig2:
            img2 = save_fig_to_tempfile(fig2)
            doc.add_paragraph("Gold vs Signal")
            doc.add_picture(img2, width=Inches(6))

        # contributions legend depends on global `labels`, so temporarily swap it
        global labels
        labels_backup = labels
        labels = labels_dict
        fig3 = build_contrib_figure(res_panel)
        labels = labels_backup

        if fig3:
            img3 = save_fig_to_tempfile(fig3)
            doc.add_paragraph("Contributions to Gold Signal (stacked)")
            doc.add_picture(img3, width=Inches(6))

    add_mode_section(left_title, resL, coreL, thrL, wL, labels_left)
    doc.add_page_break()
    add_mode_section(right_title, resR, coreR, thrR, wR, labels_right)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def build_word_report(latest_row: pd.Series,
                      prev_row: pd.Series,
                      thresholds: dict,
                      weights: dict,
                      trigger_info: dict,
                      labels: dict,
                      legend_map: dict,
                      core_keys: list,
                      mode: str,
                      crisis_year: str,
                      res_plot: pd.DataFrame,
                      include_indicator_charts=True,
                      gold_stats = None) -> bytes:
    doc = Document()

    def add_colored_regime_line(doc, regime_label: str):
        # Simple color palette
        colors = {
            "Structural Bull": RGBColor(0x00, 0x80, 0x00),  # green
            "Positive": RGBColor(0x00, 0x80, 0x00),  # green
            "Neutral": RGBColor(0x80, 0x80, 0x00),  # olive
            "Vulnerable": RGBColor(0xC0, 0x60, 0x00),  # orange
            "Structural Headwind": RGBColor(0xB0, 0x00, 0x00),  # red
        }

        p = doc.add_paragraph()
        p.add_run("Regime: ").bold = True
        r = p.add_run(regime_label)
        r.bold = True
        r.font.color.rgb = colors.get(regime_label, RGBColor(0x00, 0x00, 0x00))

    doc.add_heading("Gold Macro Cockpit — Monthly Assessment", level=1)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Date (latest data point): {latest_row.name.date()}")

    if mode == "Structural Regime (today)":
        doc.add_paragraph("Mode: Structural Regime (today)")
    elif mode == "Crisis Similarity (template)":
        doc.add_paragraph(f"Mode: Crisis Similarity ({crisis_year} template)")
    else:
        doc.add_paragraph("Mode: Market Acceleration (fast)")

    # Summary
    doc.add_heading("Summary", level=2)

    if mode == "Structural Regime (today)":
        doc.add_paragraph(
            "Model Mode: Structural Regime — evaluates whether today’s economic conditions "
            "are structurally supportive or hostile to gold using fixed economic thresholds."
        )
        doc.add_paragraph(
            "Regime Interpretation: "
            "Structural Bull = strong macro tailwind/upward push to gold; "
            "Positive = Mild tailwind; "
            "Neutral = balanced; "
            "Vulnerable = mild headwind/downward push to gold; "
            "Structural Headwind = strong macro pressure/downward push to gold."
        )
    elif mode == "Crisis Similarity (template)":
        doc.add_paragraph(
            f"Model Mode: Crisis Similarity ({crisis_year} template) — compares current "
            "macro conditions to the selected historical crisis window using quantile thresholds."
        )
        doc.add_paragraph(
            "Regime Interpretation: Structural Bull = highly similar to crisis-style "
            "gold-supportive regime; Vulnerable/Headwind = materially different from crisis regime."
        )
        doc.add_paragraph(
            "Inflation note (Crisis Similarity): If inflation looks “bearish” here, it usually just means "
            "inflation today is much lower than it was in that crisis (e.g., 1980). It does not automatically mean gold will fall."
        )

        if gold_stats and gold_stats.get("6m", {}).get("ok"):
            g6 = gold_stats["6m"]
            doc.add_paragraph(
                f"Crisis-conditioned gold response (within {crisis_year} window): "
                f"Among months with similar macro signal (±0.15), gold’s forward {g6['horizon_m']}M return "
                f"had mean {g6['mean_ret']:.1f}%, median {g6['median_ret']:.1f}%, "
                f"and P(>0) {100 * g6['p_pos']:.0f}% (n={g6['n_samples']})."
            )
        else:
            doc.add_paragraph(
                "Crisis-conditioned gold response: unavailable (insufficient similar months or missing GOLD_USD).")

    else:
        doc.add_paragraph("Market Acceleration mode answers: “Is the market currently rewarding gold (or pressuring it) regardless of the slow macro regime?"
        )

    # RSI overlay (non-scored diagnostic)
    if res_plot is not None and "RSI_14" in res_plot.columns:
        rsi_last = float(res_plot["RSI_14"].dropna().iloc[-1]) if res_plot["RSI_14"].dropna().shape[0] else np.nan
        rsi_z_last = float(res_plot["RSI_Z_60"].dropna().iloc[-1]) if (
                    "RSI_Z_60" in res_plot.columns and res_plot["RSI_Z_60"].dropna().shape[0]) else np.nan
        rsi_slope_last = float(res_plot["RSI_SLOPE_3M"].dropna().iloc[-1]) if (
                    "RSI_SLOPE_3M" in res_plot.columns and res_plot["RSI_SLOPE_3M"].dropna().shape[0]) else np.nan

        doc.add_paragraph(
            "RSI overlay (not scored): "
            f"RSI(14M)={('—' if pd.isna(rsi_last) else f'{rsi_last:.0f}')}, "
            f"RSI z-score (5Y)={('—' if pd.isna(rsi_z_last) else f'{rsi_z_last:+.2f}')}, "
            f"RSI slope (3M)={('—' if pd.isna(rsi_slope_last) else f'{rsi_slope_last:+.0f}')}. "
            f"Interpretation: {classify_rsi_overlay(rsi_last, rsi_z_last, rsi_slope_last)}"
        )

    if mode == "Crisis Similarity (template)":
        doc.add_paragraph("Note: In Crisis Similarity mode Deficit data are shown for context but not included in the score.")

    doc.add_paragraph(
        "Signal Definition: Each indicator is scored (-1 / 0 / +1) relative to its threshold band "
        "and combined using normalized weights."
    )

    doc.add_paragraph(
        "Note: Regime classification describes macro backdrop (financial) conditions, not direct trading signals."
    )

    doc.add_paragraph("")  # spacing

    sig = float(latest_row["SIGNAL"])
    sig_prev = float(prev_row["SIGNAL"])
    doc.add_paragraph(f"Weighted Signal: {sig:.2f} (Δ {sig - sig_prev:+.2f} vs previous month)")

    regime_label = str(latest_row.get("REGIME", "—"))
    add_colored_regime_line(doc, regime_label)

    def simple_narrative(latest_row: pd.Series, core_keys: list, labels: dict, regime_label: str) -> str:
        # Get contributions
        contribs = []
        for k in core_keys:
            c = latest_row.get(f"{k}_CONTRIB", 0.0)
            if pd.notna(c) and abs(float(c)) > 1e-9:
                contribs.append((k, float(c)))

        # Sort by absolute impact
        contribs.sort(key=lambda x: abs(x[1]), reverse=True)

        # Pick top positive and top negative drivers (if any)
        pos = next(((k, c) for k, c in contribs if c > 0), None)
        neg = next(((k, c) for k, c in contribs if c < 0), None)

        # Friendly short regime phrase
        regime_simple = {
            "Structural Bull": "strongly supportive",
            "Positive": "somewhat supportive",
            "Neutral": "mixed",
            "Vulnerable": "somewhat negative",
            "Structural Headwind": "negative",
        }.get(regime_label, "mixed")

        # Driver names (keep very simple)
        def short_name(k: str) -> str:
            # you can customize these if you want even simpler
            m = {
                "REAL_YIELD_CPI": "real yields",
                "CPI_YOY": "inflation",
                "USD_12M_CHG": "the dollar",
                "CURVE_10Y_3M": "the yield curve",
                "DEFICIT_GDP": "the deficit",
                "REAL_YIELD_TIPS10": "TIPS real yields",
                "HY_OAS": "credit stress",
            }
            return m.get(k, labels.get(k, k))

        if pos and neg:
            return f"Gold outlook is {regime_simple}, helped by {short_name(pos[0])} but hurt by {short_name(neg[0])}."
        if pos:
            return f"Gold outlook is {regime_simple}, mainly helped by {short_name(pos[0])}."
        if neg:
            return f"Gold outlook is {regime_simple}, mainly hurt by {short_name(neg[0])}."
        return f"Gold outlook is {regime_simple} based on the current macro mix."

    narr = simple_narrative(latest_row, core_keys, labels, regime_label)
    doc.add_paragraph(f"Plain-language summary: {narr}")

    doc.add_paragraph(
        f"Weights: {', '.join([f'{k}={weights[k]:.2f}' for k in core_keys])}"
    )

    # Trigger status
    doc.add_heading("Triggers", level=2)

    doc.add_paragraph(
        "Trigger Logic: A Bull trigger activates when the weighted signal remains "
        f"≥ {trigger_info['trig_hi']:.2f} for {trigger_info['persist']} consecutive months. "
        "A Bear trigger activates when the signal remains below the lower threshold "
        f"for the same persistence period."
    )

    doc.add_paragraph(f"Bull trigger: {'ON' if trigger_info['bull_now'] else 'OFF'} "
                      f"(threshold {trigger_info['trig_hi']:.2f}, persistence {trigger_info['persist']} months)")
    doc.add_paragraph(f"Bear trigger: {'ON' if trigger_info['bear_now'] else 'OFF'} "
                      f"(threshold {trigger_info['trig_lo']:.2f}, persistence {trigger_info['persist']} months)")

    # Weights
    doc.add_heading("Model Weights", level=2)
    t = doc.add_table(rows=1, cols=2)
    hdr = t.rows[0].cells
    hdr[0].text = "Indicator"
    hdr[1].text = "Weight"
    for k, w in weights.items():
        row = t.add_row().cells
        row[0].text = legend_map.get(k, k)
        row[1].text = f"{w:.2f}"

    # Indicator states table
    doc.add_heading("Indicator States (latest)", level=2)
    table = doc.add_table(rows=1, cols=5)
    h = table.rows[0].cells
    h[0].text = "Indicator"
    h[1].text = "Latest"
    h[2].text = "State (-1: bearish/0: neutral/+1: bullish)"
    h[3].text = "Contribution"
    h[4].text = "Thresholds (bull/bear)"

    # ----------------------------
    # Indicator charts (below indicator states)
    # ----------------------------

    core = core_keys
    for col in core:
        row = table.add_row().cells
        row[0].text = labels.get(col, col)
        row[1].text = f"{float(latest_row[col]):.2f}"
        row[2].text = str(int(latest_row[col + "_STATE"]))
        row[3].text = f"{float(latest_row[col + '_CONTRIB']):+.2f}"
        row[4].text = f"{thresholds[col]['bull']:.2f} / {thresholds[col]['bear']:.2f}"

    doc.add_paragraph("")
    if mode == "Structural Regime (today)":
        doc.add_paragraph("Interpretation note: States are computed from fixed structural thresholds "
                          "and combined using the weighted signal.")
    elif mode == "Crisis Similarity (template)":
        doc.add_paragraph("Interpretation note: States are computed from crisis-window-derived thresholds "
                          "(quantiles) and combined using the weighted signal.")
    else:
        doc.add_paragraph("Interpretation note: It uses fast-moving features (Gold momentum, real-yield velocity, stress velocity, USD velocity) scored with fixed/quantile thresholds "
                          "to produce an acceleration signal intended to capture rallies/drawdowns that can lead macro regime shifts."
                        )

        # --- Indicator charts (below Indicator States)
        if include_indicator_charts:
            doc.add_heading("Indicator Charts (history)", level=2)

            for k in core_keys:
                fig_ind = build_indicator_history_figure(
                    res_plot=res_plot,
                    key=k,
                    label=labels.get(k, k),
                    thresholds=thresholds,
                    mode=mode
                )
                if fig_ind:
                    imgp = save_fig_to_tempfile(fig_ind)
                    doc.add_picture(imgp, width=Inches(6))

        doc.add_heading("Charts", level=2)

        fig1 = build_signal_figure(res_plot)
        img1 = save_fig_to_tempfile(fig1)
        doc.add_paragraph("Signal over time")
        doc.add_picture(img1, width=Inches(6))

        fig2 = build_gold_signal_figure(res_plot)
        if fig2:
            img2 = save_fig_to_tempfile(fig2)
            doc.add_paragraph("Gold vs Signal")
            doc.add_picture(img2, width=Inches(6))

        fig3 = build_contrib_figure(res_plot)
        if fig3:
            img3 = save_fig_to_tempfile(fig3)
            doc.add_paragraph("Contributions (stacked)")
            doc.add_picture(img3, width=Inches(6))

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()




    def build_indicator_history_figure(res_plot: pd.DataFrame, key: str, label: str, thresholds: dict, mode: str):
        if key not in res_plot.columns:
            return None

        s = res_plot[key].dropna()
        if s.empty:
            return None

        fig, ax = plt.subplots(figsize=(8, 2.8))
        ax.plot(s.index, s.values, label=label)

        # Threshold lines (bull/bear) if available
        if thresholds and key in thresholds:
            add_indicator_threshold_lines(ax, key, thresholds, mode)

        ax.set_title(label)
        ax.set_xlabel("Date")
        ax.set_ylabel(label)
        ax.legend(fontsize=8)
        ax.grid(True, alpha=0.35)
        return fig

# -----------------------------
# Streamlit UI (single tab)
# -----------------------------

st.set_page_config(page_title="Gold Macro Cockpit", layout="wide")


def simple_explainer(title: str, body: str):
    with st.expander(title):
        st.markdown(body)


st.title("MRMI - Macro Regime & Market Intelligence platform")

with st.expander("How to interpret this model"):
    st.markdown("""
**Model Modes**

- **Structural Regime (today):** Evaluates whether the current macro backdrop (economic environment) is structurally supportive or hostile to gold using fixed economic thresholds.
- **Crisis Similarity (template):** Measures how similar today’s macro conditions are to a selected historical crisis regime.

**Regime Labels**

- 🟢 Structural Bull → Strong macro tailwind / Strong trends pushing gold higher
- 🟢 Positive → Mild tailwind / Mild trends pushing gold higher
- 🟡 Neutral → Balanced backdrop / Balanced economic environment
- 🔴 Vulnerable → Mild macro headwind / Mild trends pushing gold lower
- 🔴 Structural Headwind → Strong macro pressure / Strong trends pushing gold lower

**Triggers**

Bull trigger activates when the signal stays above the upper threshold for the selected persistence period.  
Bear trigger activates when it stays below the lower threshold.

Note: This is a macro regime model, not a short-term trading signal.
      In similarity mode, bearish means ‘dissimilar to template’.
""")

@st.cache_data(ttl=3600, show_spinner=True)
def load_all(monthly_method: str) -> pd.DataFrame:
    return build_features(monthly_method=monthly_method)

with st.sidebar:
    st.header("Settings")
    simple_explainer("What this section means in simple words", """
This is where you choose **how the dashboard thinks and what it shows**.

Here you select the model mode, weights, thresholds, chart windows, and optional fast indicators. In simple terms, this section lets you tune the research terminal to match the question you want to ask.
    """)
    theme = theme_picker(key="theme_picker_global")

apply_theme(theme)

with st.sidebar:

    # ----------------------------
    # 1) Mode + crisis window
    # ----------------------------
    mode = st.selectbox(
        "Model mode",
        ["Structural Regime (today)", "Crisis Similarity (template)", "Market Acceleration (fast)"],
        index=2
    )

    # ----------------------------
    # Safety defaults (must exist for report/export even if UI/data blocks don't run)
    # ----------------------------
    trig_hi = 0.60
    trig_lo = -0.60
    persist = 2
    bull_now = False
    bear_now = False

    enable_compare = st.checkbox(
        "Enable Comparison View",
        value=False,
        help="Compare any two modes side-by-side. Uses current slider settings (or defaults if untouched)."
    )

    compare_modes = None
    if enable_compare:
        compare_modes = st.multiselect(
            "Compare exactly two modes",
            ["Structural Regime (today)", "Crisis Similarity (template)", "Market Acceleration (fast)"],
            default=["Structural Regime (today)", "Market Acceleration (fast)"],
            max_selections=2
        )
    # ----------------------------
    # Crisis preset (shown when Crisis mode is selected OR when comparison includes Crisis)
    # ----------------------------
    need_crisis_ui = (mode == "Crisis Similarity (template)") or (
                enable_compare and compare_modes and ("Crisis Similarity (template)" in compare_modes))

    crisis = None
    win_start, win_end = None, None

    if need_crisis_ui:
        st.subheader("Crisis Threshold Package")
        crisis = st.selectbox(
            "Preset window",
            [
                "1929 Great Depression",
                "1974 Oil Shock",
                "1980 Volcker Shock",
                "2011 Euro Crisis",
                "2020 Pandemic",
                "Custom",
            ],
            index=2
        )

        preset_windows = {
            # Note: many modern series won't exist; the model will use whichever indicators have enough data.
            "1929 Great Depression": ("1929-01-01", "1933-12-31"),
            "1974 Oil Shock": ("1973-01-01", "1975-12-31"),
            "1980 Volcker Shock": ("1979-01-01", "1982-12-31"),
            "2011 Euro Crisis": ("2010-01-01", "2012-12-31"),
            "2020 Pandemic": ("2019-01-01", "2021-12-31"),
        }

        if crisis == "Custom":
            start = st.date_input("Window start", value=pd.to_datetime("2019-01-01"))
            end = st.date_input("Window end", value=pd.to_datetime("2021-12-31"))
            win_start, win_end = str(start), str(end)
        else:
            win_start, win_end = preset_windows[crisis]

        st.caption(f"Window: {win_start} → {win_end}")
        st.info(
            "Crisis Similarity mode: Deficit % GDP is shown for context, but it is NOT included in the score "
            "(to avoid low-frequency artifacts)."
        )

    # ----------------------------
    # 2) Weights + optional extras
    # ----------------------------

    show_indicator_thresholds = st.checkbox("Show thresholds on indicator charts", value=True)

    st.subheader("Weights (auto-normalized)")
    w_real = st.slider("Real yield", 0.0, 1.0, 0.35, 0.01)
    w_usd = st.slider("Dollar", 0.0, 1.0, 0.25, 0.01)
    w_infl = st.slider("Inflation", 0.0, 1.0, 0.20, 0.01)
    w_curve = st.slider("Curve", 0.0, 1.0, 0.10, 0.01)
    w_fisc = st.slider("Fiscal dominance", 0.0, 1.0, 0.10, 0.01)

    st.subheader("Additional Structural Indicators")
    include_tips = st.checkbox(
        "Include 10Y TIPS real yield (DFII10)",
        value=False,
        disabled=(mode != "Structural Regime (today)")
    )
    w_tips = st.slider(
        "Weight: TIPS real yield",
        0.0, 1.0, 0.10, 0.01,
        disabled=(not include_tips or mode != "Structural Regime (today)")
    )

    include_hy = st.checkbox(
        "Include HY credit spread in Gold Structural model (HY OAS)",
        value=False,
        disabled=(mode != "Structural Regime (today)")
    )
    w_hy = st.slider(
        "Include HY credit spread in Gold Structural model (HY OAS)",
        0.0, 1.0, 0.10, 0.01,
        disabled=(not include_hy or mode != "Structural Regime (today)")
    )

    # In crisis similarity mode, deficit must not be scored
    if mode != "Structural Regime (today)":
        w_fisc = 0.0

    # ----------------------------
    # 3) Monthly aggregation + data load
    # ----------------------------
    st.subheader("Monthly aggregation")
    monthly_method = st.selectbox("Method for daily series", ["avg", "eom"], index=0)

    # ----------------------------
    # Defaults so intraday vars ALWAYS exist
    # (must be defined BEFORE any reference)
    # ----------------------------
    show_intraday_rsi = False
    intraday_interval = "15m"
    intraday_lookback_days = 10

    df = load_all(monthly_method)
    res_base = df.copy()

    # ----------------------------
    # NASDAQ Screener tickers (used by Tab 2)
    # ----------------------------
    st.subheader("Intraday RSI Screener (Top 10)")

    _default_top10 = "AAPL MSFT NVDA AMZN META GOOGL TSLA AVGO COST AMD WMT ASML MU NFLX PLTR CSCO AMAT LRCX PEP INTC "

    MAX_TICKERS = 20

    raw_top10 = st.text_area(
        f"Tickers (space/comma/newline separated) — max {MAX_TICKERS}",
        value=_default_top10,
        help="These tickers will be used in the Intraday RSI Screener tab."
    )

    tickers_top10 = (
        pd.Series(raw_top10.replace(",", " ").split())
        .astype(str).str.strip().str.upper()
        .tolist()
    )
    tickers_top10 = [t for t in tickers_top10 if t]
    tickers_top10 = list(dict.fromkeys(tickers_top10))[:MAX_TICKERS]

    # Build acceleration features (used by Mode C and/or Combined View)
    res_accel_base = compute_accel_features(res_base)

    # ----------------------------
    # Build STRUCTURAL weights (always build, so Structural results are consistent even in comparison)
    # ----------------------------
    raw_struct = {
        "REAL_YIELD_CPI": w_real,
        "CPI_YOY": w_infl,
        "USD_12M_CHG": w_usd,
        "CURVE_10Y_3M": w_curve,
        "DEFICIT_GDP": w_fisc,
        "REAL_YIELD_TIPS10": (w_tips if include_tips else 0.0),
        "HY_OAS": (w_hy if include_hy else 0.0),
    }

    # Avoid double counting real yields if TIPS is enabled + exists
    if include_tips and ("REAL_YIELD_TIPS10" in res_base.columns) and (res_base["REAL_YIELD_TIPS10"].dropna().size > 0):
        raw_struct["REAL_YIELD_CPI"] = 0.0

    active = {k: v for k, v in raw_struct.items() if v > 0}
    tot = sum(active.values()) or 1.0
    weights_struct = {k: v / tot for k, v in active.items()}

    # ----------------------------
    # Build CRISIS weights (same sliders but deficit excluded; structural-only extras excluded)
    # ----------------------------
    raw_crisis = {
        "REAL_YIELD_CPI": w_real,
        "CPI_YOY": w_infl,
        "USD_12M_CHG": w_usd,
        "CURVE_10Y_3M": w_curve,
        "DEFICIT_GDP": 0.0,  # excluded in similarity scoring
        "REAL_YIELD_TIPS10": 0.0,  # optional structural-only
        "HY_OAS": 0.0,  # optional structural-only
    }
    active = {k: v for k, v in raw_crisis.items() if v > 0}
    tot = sum(active.values()) or 1.0
    weights_crisis = {k: v / tot for k, v in active.items()}

    # For the "single run" mode output, keep a convenience alias called `weights`
    if mode == "Structural Regime (today)":
        weights = weights_struct
    elif mode == "Crisis Similarity (template)":
        weights = weights_crisis
    else:
        # Mode C uses its own weights inside run_accel()
        weights = weights_struct  # used only for display/other non-scoring UI

    # ----------------------------
    # 4) Chart controls
    # ----------------------------
    st.subheader("Charts")
    simple_explainer("What this section means in simple words", """
These settings control **how much history you see on the screen**.

- **Indicators' Plot Timeframe** changes how many months are shown in the charts.
- **History view** changes the historical window, for example full history, recent years, or a crisis window.

This section does **not** change the raw data itself. It mainly changes **the window through which you look at the data**.
    """)
    lookback_options = [6, 12, 24, 36, 60, 120, 180, 240, 360]
    lookback_label_options = [str(x) for x in lookback_options] + ["All available"]

    lookback_label = st.selectbox(
        "Indicators' Plot Timeframe (months)",
        lookback_label_options,
        index=lookback_label_options.index("60") if "60" in lookback_label_options else 0
    )

    if lookback_label == "All available":
        lookback = None
    else:
        lookback = int(lookback_label)
    history_view = st.selectbox(
        "History view",
        ["Full history", "Last 15y", "Last 5y", "Crisis window only", "Crisis window ±5y"],
        index=0
    )

    # ----------------------------
    # 5) Trigger controls
    # ----------------------------
    st.subheader("Trigger")
    simple_explainer("What this section means in simple words", """
These controls define the **line in the sand** for strong bullish or bearish signals.

- The **bull trigger** says when the score is strong enough to count as clearly positive.
- The **bear trigger** says when the score is weak enough to count as clearly negative.
- **Persistence months** asks whether the signal should stay strong for more than one month before you trust it.

Think of this as the platform's **confidence filter**.
    """)
    trig_hi = st.slider("Bull trigger (+)", 0.0, 1.5, 0.60, 0.05)
    trig_lo = st.slider("Bear trigger (−)", -1.5, 0.0, -0.60, 0.05)
    persist = st.slider("Persistence months", 1, 3, 2, 1)

    # ----------------------------
    # Monte Carlo Settings
    # ----------------------------
    st.subheader("Monte Carlo Settings")
    simple_explainer("What this section means in simple words", """
These settings control the **possible future path analysis**.

- **Horizon** tells the model how far into the future to simulate.
- **Simulations** tells it how many alternative paths to generate.

More simulations usually make the result smoother, but they can take longer.
    """)

    mc_h = st.slider(
        "Monte Carlo horizon (months)",
        min_value=3,
        max_value=24,
        value=12,
        step=1
    )

    mc_n = st.slider(
        "Monte Carlo simulations",
        min_value=500,
        max_value=5000,
        value=2000,
        step=500
    )

    # ----------------------------
    # Market Acceleration controls (Mode C) — shown when Accel mode is selected OR comparison includes Accel
    # ----------------------------

    st.subheader("Live updates")
    simple_explainer("What this section means in simple words", """
This section controls whether the dashboard refreshes itself automatically.

Use it when you want the screen to keep updating during the trading day. If you leave it off, the platform only refreshes when you do it manually.
    """)

    live_updates = st.checkbox(
        "Enable live updates (auto-refresh)",
        value=False,
        help="Auto-refresh the app on a timer. Uses cached fetchers; refresh token forces updates on schedule."
    )

    refresh_seconds = st.selectbox(
        "Refresh every…",
        [60],
        index=0,
        disabled=not live_updates
    )

    # Token changes every refresh interval; use it to bust cache on schedule
    refresh_token = int(time.time() // int(refresh_seconds)) if live_updates else 0

    # Trigger refresh for the whole page
    auto_refresh(refresh_seconds, live_updates)

    need_accel_ui = (mode == "Market Acceleration (fast)") or (
                enable_compare and compare_modes and ("Market Acceleration (fast)" in compare_modes))

    # Defaults (so variables exist even if UI is hidden)
    accel_method = "Fixed (recommended)"
    w_g, w_ry, w_st, w_u = 0.35, 0.25, 0.25, 0.15

    if need_accel_ui:
        st.subheader("Acceleration Settings")
        simple_explainer("What this section means in simple words", """
This section controls the **fast-moving part** of the model.

It looks at whether gold is gaining or losing momentum right now by checking things like recent gold performance, changes in real yields, changes in stress, and changes in the dollar.

In simple terms: this is the part that asks whether the market is **speeding up** or **slowing down** for gold.
        """)

        accel_method = st.selectbox(
            "Acceleration thresholds",
            ["Fixed (recommended)", "Quantiles (2000-present)"],
            index=0
        )

        st.caption("Weights (auto-normalized)")
        w_g = st.slider("Gold momentum (6M return)", 0.0, 1.0, 0.35, 0.01)
        w_ry = st.slider("Real yield velocity (3M change)", 0.0, 1.0, 0.25, 0.01)
        w_st = st.slider("Stress velocity (3M change)", 0.0, 1.0, 0.25, 0.01)
        w_u = st.slider("USD velocity (3M change)", 0.0, 1.0, 0.15, 0.01)

        # ----------------------------
        # Intraday RSI controls (Acceleration only)
        # ----------------------------
        st.subheader("Intraday RSI (optional)")
        simple_explainer("What this section means in simple words", """
This section adds a **short-term timing view**.

It uses intraday RSI to show whether gold looks stretched upward, stretched downward, or relatively balanced over short timeframes such as minutes or hours.

This is mainly useful for timing and monitoring, not for the slow macro regime itself.
        """)

        show_intraday_rsi = st.checkbox(
            "Show intraday RSI",
            value=True,
            help="Fetch intraday gold (yfinance) and compute RSI(14) on 5m/15m bars. Best-effort; may be unavailable on weekends/holidays."
        )

        intraday_interval = st.selectbox(
            "Intraday interval",
            ["1m", "5m", "15m", "30m", "60m"],
            index=0,
            disabled=(not show_intraday_rsi),
            key="intraday_interval",
        )

        # ----------------------------
        # Intraday lookback options depend on interval (yfinance limits)
        # ----------------------------
        LOOKBACK_BY_INTERVAL = {
            "1m": [1, 2, 3],
            "2m": [1, 2, 3, 5, 7, 10, 15, 30, 60],
            "5m": [1, 2, 3, 5, 7, 10, 15, 30, 60],
            "15m": [1, 2, 3, 5, 7, 10, 15, 30, 60],
            "30m": [1, 2, 3, 5, 7, 10, 15, 30, 60],
            "60m": [5, 7, 10, 15, 30, 60, 90, 180, 365, 730],
        }

        valid_lookbacks = LOOKBACK_BY_INTERVAL.get(intraday_interval, [1, 2, 3, 5, 7, 10, 15, 30, 60])

        # Preserve previous selection if still valid; else clamp to max valid
        prev = st.session_state.get("intraday_lookback_days", None)
        default_val = prev if (prev in valid_lookbacks) else valid_lookbacks[-1]

        intraday_lookback_days = st.selectbox(
            "Intraday lookback (days)",
            valid_lookbacks,
            index=valid_lookbacks.index(default_val),
            disabled=(not show_intraday_rsi),
            key="intraday_lookback_days"
        )

        # ----------------------------
        # Intraday data (Acceleration only) — fetch once, store in attrs
        # ----------------------------
        if mode == "Market Acceleration (fast)" and show_intraday_rsi:
            # Best-effort ticker list (intraday availability varies)
            intra_tickers = ["GC=F", "XAUUSD=X", "GLD"]

            gold_intra = yf_intraday_close(
                tickers=["GC=F", "XAUUSD=X", "GLD"],
                interval=intraday_interval,
                lookback_days=intraday_lookback_days,
                refresh_token=refresh_token
            )

            if isinstance(gold_intra, pd.Series) and not gold_intra.dropna().empty:
                res_base.attrs["gold_intraday"] = gold_intra

                pack = compute_intraday_rsi_pack(gold_intra, rsi_period=14)
                if pack:
                    # Store the core RSI intraday
                    res_base.attrs["RSI_14_INTRA"] = pack.get("RSI_14_INTRA")
                    # Optional extras (safe to keep even if you don’t use them in UI yet)
                    res_base.attrs["RSI_14_INTRA_Z_2W"] = pack.get("RSI_14_INTRA_Z_2W")
                    res_base.attrs["RSI_14_INTRA_SLOPE_3H"] = pack.get("RSI_14_INTRA_SLOPE_3H")
            else:
                # Store empty marker so we don't try again within the same run
                res_base.attrs["gold_intraday"] = pd.Series(dtype=float)
                res_base.attrs["RSI_14_INTRA"] = pd.Series(dtype=float)

    # ----------------------------
    # 6) Threshold Matrix placeholder (filled later)
    # ----------------------------
    st.markdown("---")
    st.subheader("Threshold Matrix")
    simple_explainer("What this section means in simple words", """
This matrix shows the **rules behind the score**.

For each indicator, it shows what the platform treats as bullish, bearish, and in-between. In other words, this is where you can see **why** the model classified a value as supportive or unsupportive.
    """)
    threshold_box = st.container()

    # ----------------------------
    # 7) Refresh
    # ----------------------------
    st.markdown("---")
    if st.button("Refresh data"):
        st.cache_data.clear()
        st.rerun()


# Mode 1 corrected working frame

dirs_struct = {
    "REAL_YIELD_CPI": False,   # lower real yield bullish
    "CPI_YOY": True,           # higher inflation bullish
    "USD_12M_CHG": False,      # higher USD change bearish for gold (so lower is bullish)
    "CURVE_10Y_3M": False,     # more inverted bullish
    "DEFICIT_GDP": False,       # more negative deficit bullish
    "REAL_YIELD_TIPS10": False,  # lower real yield bullish
    "HY_OAS": True,  # higher stress bullish
}

dirs_market = {
    "REAL_YIELD_CPI": False,
    "USD_12M_CHG": False,
    "CURVE_10Y_3M": True,
    "HY_OAS": False,
    "QQQ_ABOVE_MA200": True,
    "QQQ_MA50_SLOPE_20D": True,
    "MARKET_BREADTH_ABOVE_MA200": True,
}

labels_market = {
    "REAL_YIELD_CPI": "Real yield (10Y − CPI YoY)",
    "USD_12M_CHG": "USD 12M % change (TWEX)",
    "CURVE_10Y_3M": "Curve (10Y–3M)",
    "HY_OAS": "High Yield OAS (20D MA)",
    "QQQ_ABOVE_MA200": "QQQ above MA200",
    "QQQ_MA50_SLOPE_20D": "QQQ MA50 slope (20d, %)",
    "MARKET_BREADTH_ABOVE_MA200": "Market breadth (% above MA200)",
}
# Crisis similarity: deficit is not used in this mode in order to avoid low-frequency quantile artifacts
dirs_crisis = {
    "REAL_YIELD_CPI": False,
    "CPI_YOY": True,
    "USD_12M_CHG": False,
    "CURVE_10Y_3M": False,
    "REAL_YIELD_TIPS10": False,
    "HY_OAS": True,
}

labels = {
"REAL_YIELD_CPI": "Real yield (10Y − CPI YoY)",
"CPI_YOY": "Inflation (CPI YoY)",
"USD_12M_CHG": "USD 12M % change (TWEX)",
"CURVE_10Y_3M": "Curve (10Y–3M)",
"DEFICIT_GDP": "Deficit % GDP",
"REAL_YIELD_TIPS10": "Real yield (10Y TIPS, DFII10, 20D MA)",
"HY_OAS": "High Yield OAS (20D MA)",
"QQQ_ABOVE_MA200": "QQQ above MA200",
"QQQ_MA50_SLOPE_20D": "QQQ MA50 slope (20d, %)",
"MARKET_BREADTH_ABOVE_MA200": "Market breadth (% above MA200)",
}

dirs = dirs_struct if mode == "Structural Regime (today)" else dirs_crisis

dirs_accel = {
    "GOLD_6M_RET": True,        # higher return bullish
    "REALYIELD_3M_CHG": False,  # falling real yields bullish
    "STRESS_3M_CHG": True,      # rising stress bullish
    "USD_3M_CHG": False,        # falling USD bullish
}

labels_accel = {
    "GOLD_6M_RET": "Gold 6M return (%)",
    "REALYIELD_3M_CHG": "Real yield 3M change (pp, TIPS preferred)",
    "STRESS_3M_CHG": "Stress 3M change (pp, HY OAS)",
    "USD_3M_CHG": "USD 3M change (%) (TWEX)",
}

def bucket_accel(x: float) -> str:
    if pd.isna(x):
        return "No data"
    if x >= 0.60:
        return "Acceleration Bull"
    if x >= 0.20:
        return "Positive"
    if x > -0.20:
        return "Neutral"
    if x > -0.60:
        return "Vulnerable"
    return "Acceleration Headwind"

def run_structural(res_base: pd.DataFrame, weights_struct: dict, dirs_struct: dict) -> tuple[pd.DataFrame, list, dict, dict, dict]:
    thresholds = structural_thresholds()
    candidate_keys = [k for k in dirs_struct.keys() if k in weights_struct]
    core_keys = [k for k in candidate_keys if (k in thresholds) and (k in res_base.columns) and (res_base[k].dropna().size > 0)]

    dirs = {k: dirs_struct[k] for k in core_keys}
    weights = {k: weights_struct[k] for k in core_keys}
    tot = sum(weights.values()) or 1.0
    weights = {k: v / tot for k, v in weights.items()}
    thresholds = {k: thresholds[k] for k in core_keys}

    res = compute_signal(res_base, thresholds, weights, dirs)

    if all((c + "_CONTRIB") in res.columns for c in core_keys):
        contrib_sum = res[[c + "_CONTRIB" for c in core_keys]].sum(axis=1)
        max_diff = (contrib_sum - res["SIGNAL"]).abs().max()
        print("Max SIGNAL vs sum(CONTRIB) diff:", max_diff)

    print("Mode:", mode)
    print("Core keys:", core_keys)
    print("Latest signal:", float(res['SIGNAL'].dropna().iloc[-1]))
    print("Latest regime:", str(res['REGIME'].dropna().iloc[-1]))

    return res, core_keys, thresholds, dirs, weights

def run_market_structural(res_base: pd.DataFrame) -> tuple[pd.DataFrame, list, dict, dict, dict]:
    thresholds = market_structural_thresholds()

    # simple fixed weights for v1
    weights_market = {
        "REAL_YIELD_CPI": 0.18,
        "USD_12M_CHG": 0.12,
        "CURVE_10Y_3M": 0.12,
        "HY_OAS": 0.20,
        "QQQ_ABOVE_MA200": 0.14,
        "QQQ_MA50_SLOPE_20D": 0.09,
        "MARKET_BREADTH_ABOVE_MA200": 0.15,
    }

    candidate_keys = [k for k in dirs_market.keys() if k in thresholds]
    core_keys = [
        k for k in candidate_keys
        if (k in res_base.columns) and (res_base[k].dropna().size > 0)
    ]

    dirs = {k: dirs_market[k] for k in core_keys}
    weights = {k: weights_market[k] for k in core_keys}
    tot = sum(weights.values()) or 1.0
    weights = {k: v / tot for k, v in weights.items()}
    thresholds = {k: thresholds[k] for k in core_keys}

    res = compute_signal(res_base, thresholds, weights, dirs)
    return res, core_keys, thresholds, dirs, weights

def run_crisis(res_base: pd.DataFrame, weights_crisis: dict, dirs_crisis: dict, win_start: str, win_end: str) -> tuple[pd.DataFrame, list, dict, dict, dict, dict]:
    candidate_keys = [k for k in dirs_crisis.keys() if k in weights_crisis]
    thresholds = compute_thresholds_from_window(res_base, win_start, win_end, keys=candidate_keys, dirs=dirs_crisis)

    # Graceful fail (e.g. 1929 window has no overlap with our dataset)

    if not thresholds:
        avail_start = res_base.index.min()
        avail_end = res_base.index.max()
        st.error(
            f"No data available in the selected crisis window ({win_start} → {win_end}).\n\n"
            f"Available dataset range is approximately: {avail_start.date()} → {avail_end.date()}.\n\n"
            "Pick a later crisis window (e.g., 1974/1980/2011/2020) or switch to Custom and use an overlapping range."
        )
        # Return a safe empty result pack so downstream UI/report doesn't crash
        res = res_base.copy()
        res["SIGNAL"] = np.nan
        res["REGIME"] = "No data"
        core_keys = []
        dirs = {}
        weights = {}
        gold_stats = {
            "3m": {"ok": False, "reason": "no-overlap"},
            "6m": {"ok": False, "reason": "no-overlap"},
        }
        return res, core_keys, {}, dirs, weights, gold_stats

    core_keys = list(thresholds.keys())

    # If a historical window has sparse data (e.g., 1929–1933),
    # we may end up with very few usable indicators.

    if len(core_keys) < 2:
                st.warning(
                    "Crisis window has very limited usable indicator history. "
             "Similarity score may be unstable (too few indicators available)."
            )

    dirs = {k: dirs_crisis[k] for k in core_keys}
    weights = {k: weights_crisis.get(k, 0.0) for k in core_keys}
    tot = sum(weights.values()) or 1.0
    weights = {k: v / tot for k, v in weights.items()}

    res = compute_signal(res_base, thresholds, weights, dirs)

    gold_stats_6m = crisis_conditioned_gold_stats(
        res_base, win_start, win_end, thresholds, weights, dirs,
        horizon_months=6, band=0.15
    )
    gold_stats_3m = crisis_conditioned_gold_stats(
        res_base, win_start, win_end, thresholds, weights, dirs,
        horizon_months=3, band=0.15
    )
    gold_stats = {"3m": gold_stats_3m, "6m": gold_stats_6m}

    return res, core_keys, thresholds, dirs, weights, gold_stats


def run_accel(res_accel_base: pd.DataFrame, accel_method: str, w_g: float, w_ry: float, w_st: float, w_u: float) -> tuple[pd.DataFrame, list, dict, dict, dict]:
    accel_keys = list(dirs_accel.keys())

    raw_w = {"GOLD_6M_RET": w_g, "REALYIELD_3M_CHG": w_ry, "STRESS_3M_CHG": w_st, "USD_3M_CHG": w_u}
    active = {k: v for k, v in raw_w.items() if v > 0}
    tot = sum(active.values()) or 1.0
    weights = {k: v / tot for k, v in active.items()}

    if accel_method.startswith("Fixed"):
        thresholds = accel_thresholds_fixed()
    else:
        thresholds = compute_accel_thresholds_quantiles(res_accel_base, accel_keys, dirs_accel)

    core_keys = [k for k in accel_keys if (k in thresholds and k in weights and k in res_accel_base.columns)]
    dirs = {k: dirs_accel[k] for k in core_keys}
    thresholds = {k: thresholds[k] for k in core_keys}
    weights = {k: weights[k] for k in core_keys}
    tot = sum(weights.values()) or 1.0
    weights = {k: v / tot for k, v in weights.items()}

    res = compute_signal(res_accel_base, thresholds, weights, dirs)
    # Make regime naming consistent for accel
    res["REGIME"] = res["SIGNAL"].apply(bucket_accel)
    return res, core_keys, thresholds, dirs, weights

# ----------------------------
# Mode routing (single main run)
# ----------------------------
if mode == "Structural Regime (today)":
    gold_stats = None
    res, core_keys, thresholds, dirs, weights_used = run_structural(res_base, weights_struct, dirs_struct)
    labels_used = labels
elif mode == "Crisis Similarity (template)":
    # safety: require window
    if not (win_start and win_end):
        st.error("Crisis Similarity mode requires a preset window (or Custom window).")
        st.stop()
    res, core_keys, thresholds, dirs, weights_used, gold_stats = run_crisis(res_base, weights_crisis, dirs_crisis, win_start, win_end)
    labels_used = labels
else:
    # Market Acceleration (fast)
    gold_stats = None
    res, core_keys, thresholds, dirs, weights_used = run_accel(res_accel_base, accel_method, w_g, w_ry, w_st, w_u)
    labels_used = labels_accel

# Always compute market structural regime once (used in KPI strip + Tab 3 decisions)
res_market, core_keys_market, thresholds_market, dirs_market_used, weights_market_used = run_market_structural(res_base)
latest_market = res_market.iloc[-1] if not res_market.empty else None
market_regime_now = str(latest_market.get("REGIME", "—")) if latest_market is not None else "—"
market_signal_now = float(latest_market["SIGNAL"]) if latest_market is not None and "SIGNAL" in latest_market else np.nan

# keep your downstream UI working with the expected variable names
weights = weights_used
labels = labels_used

# -------------------------------
# Sidebar Threshold Display (render into the placeholder created above)
# -------------------------------
with threshold_box:
    if mode == "Structural Regime (today)":
        title = "Structural Threshold Matrix (with direction)"
    elif mode == "Crisis Similarity (template)":
        title = "Crisis Threshold Matrix (with direction)"
    else:
        title = "Acceleration Threshold Matrix (with direction)"
    st.caption(title)

    if mode != "Structural Regime (today)":
        st.caption("Deficit % GDP is shown in cards for reference only in Crisis Similarity mode.")

    # ----------------------------
    # Grouped indicator display
    # ----------------------------
    if mode in ["Structural Regime (today)", "Crisis Similarity (template)"]:

        common_macro_keys = ["REAL_YIELD_CPI", "USD_12M_CHG", "CURVE_10Y_3M"]
        gold_only_keys = ["CPI_YOY", "DEFICIT_GDP", "REAL_YIELD_TIPS10"]
        market_only_keys = ["HY_OAS", "QQQ_ABOVE_MA200", "QQQ_MA50_SLOPE_20D", "MARKET_BREADTH_ABOVE_MA200"]


        def render_threshold_group(group_title, keys_to_show):
            shown = [k for k in keys_to_show if k in labels]
            if not shown:
                return

            st.markdown(f"#### {group_title}")

            for key in shown:
                # Gold/common side
                gold_scored = (key in core_keys) and (key in thresholds)
                # Market side
                market_scored = (key in core_keys_market) and (
                            key in thresholds_market) if core_keys_market is not None and thresholds_market is not None else False

                # Pick the correct threshold source
                if market_scored and not gold_scored:
                    th = thresholds_market[key]
                    higher_is_bull = dirs_market.get(key, False)
                elif gold_scored:
                    th = thresholds[key]
                    higher_is_bull = dirs.get(key, False)
                else:
                    continue

                bull = th["bull"]
                bear = th["bear"]
                name = labels.get(key, key)

                st.markdown(f"**{name}**")

                # Special friendly wording for binary market indicator
                if key == "QQQ_ABOVE_MA200":
                    st.caption("Bull if = Yes (above MA200)")
                    st.caption("Bear if = No (below MA200)")
                else:
                    if higher_is_bull:
                        st.caption(f"Bull if ≥ {bull:.2f}")
                        st.caption(f"Bear if ≤ {bear:.2f}")
                    else:
                        st.caption(f"Bull if ≤ {bull:.2f}")
                        st.caption(f"Bear if ≥ {bear:.2f}")

                st.markdown("---")

        render_threshold_group("Common Macro Indicators", common_macro_keys)
        render_threshold_group("Gold-Specific Indicators", gold_only_keys)
        render_threshold_group("Market-Specific Indicators", market_only_keys)

    else:
        # Acceleration mode keeps the current flat list
        for key in core_keys:
            if key not in thresholds:
                continue

            bull = thresholds[key]["bull"]
            bear = thresholds[key]["bear"]
            name = labels.get(key, key)
            higher_is_bull = dirs.get(key, False)

            st.markdown(f"**{name}**")
            if higher_is_bull:
                st.caption(f"Bull if ≥ {bull:.2f}")
                st.caption(f"Bear if ≤ {bear:.2f}")
            else:
                st.caption(f"Bull if ≤ {bull:.2f}")
                st.caption(f"Bear if ≥ {bear:.2f}")

            st.markdown("---")

# --- Top KPI strip
latest = res.iloc[-1]
prev = res.iloc[-2] if len(res) > 1 else latest
delta = float(latest["SIGNAL"] - prev["SIGNAL"])

# --- Crisis-conditioned gold response (ONLY in Crisis mode)
if mode == "Crisis Similarity (template)":
    st.subheader("Crisis-conditioned Gold Response (in-template)")
    simple_explainer("What this section means in simple words", """
This section asks a practical question: **When the world looked like this crisis template before, what did gold usually do next?**

It does not predict the future with certainty. It simply looks at similar past moments inside the chosen crisis window and summarizes how gold behaved afterward.
    """)
    if gold_stats and gold_stats.get("6m", {}).get("ok"):
        g6 = gold_stats["6m"]
        st.write(
            f"Using crisis-window months with similar macro signal (±0.15): "
            f"n={g6['n_samples']} samples | "
            f"6M fwd mean={g6['mean_ret']:.1f}% | median={g6['median_ret']:.1f}% | "
            f"P(>0)={100*g6['p_pos']:.0f}% | IQR=({g6['q25']:.1f}%, {g6['q75']:.1f}%)"
        )
    else:
        st.warning(
            f"No reliable gold-conditioned stats: "
            f"{gold_stats['6m'].get('reason') if gold_stats and '6m' in gold_stats else 'missing'}"
        )

# --- Top KPI strip
st.markdown("")  # small spacing

if mode == "Market Acceleration (fast)":

    # Row: Signal + Regime + (Monthly RSI block)  (Daily RSI block)
    rsi_col = "RSI_14_ASOF" if "RSI_14_ASOF" in res_base.columns else "RSI_14"
    rsi_last = float(res_base[rsi_col].dropna().iloc[-1]) if (
                rsi_col in res_base.columns and res_base[rsi_col].dropna().shape[0] > 0) else np.nan

    # Daily RSI overlays (aligned to month-end; may be missing)
    rsi14d_daily = res_base.attrs.get("RSI_14D_DAILY", None)
    rsi14d_z_daily = res_base.attrs.get("RSI_14D_Z_1Y_DAILY", None)
    rsi14d_slope_daily = res_base.attrs.get("RSI_14D_SLOPE_1M_DAILY", None)

    rsi_d = float(rsi14d_daily.dropna().iloc[-1]) if isinstance(rsi14d_daily,
                                                                pd.Series) and not rsi14d_daily.dropna().empty else np.nan
    rsi_d_z = float(rsi14d_z_daily.dropna().iloc[-1]) if isinstance(rsi14d_z_daily,
                                                                    pd.Series) and not rsi14d_z_daily.dropna().empty else np.nan
    rsi_d_slope = float(rsi14d_slope_daily.dropna().iloc[-1]) if isinstance(rsi14d_slope_daily,
                                                                            pd.Series) and not rsi14d_slope_daily.dropna().empty else np.nan
    # --- Acceleration KPI layout (stacked RSI style)
    k1, k2, k3, k4, k5 = st.columns([1.2, 1.2, 1, 1, 1])
    with k1:
        st.metric("Weighted Signal", f'{latest["SIGNAL"]:.2f}', f"{delta:+.2f}")
    with k2:
        st.metric("Regime", str(latest.get("REGIME", "—")))

        rsi_last = float(res_base["RSI_14"].dropna().iloc[-1]) if (
                    "RSI_14" in res_base.columns and res_base["RSI_14"].dropna().shape[0] > 0
            ) else np.nan
        rsi_z_last = float(res_base["RSI_Z_60"].dropna().iloc[-1]) if (
                    "RSI_Z_60" in res_base.columns and res_base["RSI_Z_60"].dropna().shape[0] > 0
            ) else np.nan
        rsi_slope_last = float(res_base["RSI_SLOPE_3M"].dropna().iloc[-1]) if (
                    "RSI_SLOPE_3M" in res_base.columns and res_base["RSI_SLOPE_3M"].dropna().shape[0] > 0
            ) else np.nan

      # Daily RSI overlays (aligned to month-end; may be missing)
        rsi_d = float(res_base["RSI_14D"].dropna().iloc[-1]) if (
                    "RSI_14D" in res_base.columns and res_base["RSI_14D"].dropna().shape[0] > 0
            ) else np.nan
        rsi_d_z = float(res_base["RSI_14D_Z_1Y"].dropna().iloc[-1]) if (
                    "RSI_14D_Z_1Y" in res_base.columns and res_base["RSI_14D_Z_1Y"].dropna().shape[0] > 0
            ) else np.nan
        rsi_d_slope = float(res_base["RSI_14D_SLOPE_1M"].dropna().iloc[-1]) if (
                    "RSI_14D_SLOPE_1M" in res_base.columns and res_base["RSI_14D_SLOPE_1M"].dropna().shape[0] > 0
            ) else np.nan

    with k3:
        st.markdown("**RSI (Monthly)**")
        m1, m2, m3 = st.columns([1, 1.2, 1])
        m1.metric("RSI 14M", "—" if pd.isna(rsi_last) else f"{rsi_last:.0f}")
        m2.metric("z-score (5Y)", "—" if pd.isna(rsi_z_last) else f"{rsi_z_last:+.2f}")
        m3.metric("slope (3M)", "—" if pd.isna(rsi_slope_last) else f"{rsi_slope_last:+.0f}")

        st.markdown("**RSI (Daily)**")
        d1, d2, d3 = st.columns(3)
        d1.metric("RSI 14D", "—" if pd.isna(rsi_d) else f"{rsi_d:.0f}")
        d2.metric("vs 1Y (z)", "—" if pd.isna(rsi_d_z) else f"{rsi_d_z:+.2f}")
        d3.metric("change (1M)", "—" if pd.isna(rsi_d_slope) else f"{rsi_d_slope:+.0f}")

    st.caption(f"Core data through: {str(res.index.max().date())}")

    # --- ONE merged RSI message (remove the “two similar messages” issue)
    regime_now = str(latest.get("REGIME", "—"))
    stretched_up_long = (not pd.isna(rsi_last) and rsi_last >= 80)
    stretched_up_short = (not pd.isna(rsi_d) and rsi_d >= 70)
    stretched_down_long = (not pd.isna(rsi_last) and rsi_last <= 30)
    stretched_down_short = (not pd.isna(rsi_d) and rsi_d <= 30)

    if "Bull" in regime_now or "Positive" in regime_now:
        if stretched_up_long and stretched_up_short:
            combo = "Acceleration is positive, but both RSIs are high. Momentum is strong, yet a short pause or pullback would be normal."
        elif stretched_up_long and not stretched_up_short:
            combo = "Acceleration is positive. Long-term momentum is very strong, while short-term momentum is not overheated."
        elif stretched_up_short and not stretched_up_long:
            combo = "Acceleration is positive and short-term momentum is hot. The move may be fast in the short run."
        else:
            combo = "Acceleration is positive and RSI is not extreme. That’s usually a healthy trend setup."
    elif "Headwind" in regime_now or "Vulnerable" in regime_now:
        if stretched_down_long and stretched_down_short:
            combo = "Acceleration is negative and both RSIs are very low. Selling may be exhausted, so a rebound is possible."
        elif stretched_down_short and not stretched_down_long:
            combo = "Acceleration is negative and short-term RSI is very low. That can happen near short-term bottoms."
        else:
            combo = "Acceleration is negative and RSI is not deeply oversold. Pressure may continue."
    else:
        if stretched_up_long and stretched_up_short:
            combo = "Acceleration is mixed, but RSI is high in both timeframes. This often cools off."
        elif stretched_down_long and stretched_down_short:
            combo = "Acceleration is mixed, but RSI is very low in both timeframes. This can happen near bottoms."
        else:
            combo = "Acceleration is mixed and RSI is not extreme."

    st.info(f"RSI & Acceleration: {combo}")

else:
    # Structural / Crisis KPI strip
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Gold Signal", f'{latest["SIGNAL"]:.2f}', f"{delta:.2f}")
    c2.metric("Gold Regime", str(latest.get("REGIME", "—")))
    c3.metric("Market Regime", market_regime_now)

    if mode == "Crisis Similarity (template)":
        c4.metric("Preset Window", crisis if crisis != "Custom" else f"{win_start}→{win_end}")
    else:
        c4.metric("Preset Window", "—")

    c5.metric("Core data through", str(res.index.max().date()))

import matplotlib.dates as mdates

def render_history_panel(
    title: str,
    res_panel: pd.DataFrame,
    df_panel: pd.DataFrame,
    win_start: str | None,
    win_end: str | None,
    trig_hi: float,
    trig_lo: float,
    res_market: pd.DataFrame | None = None,
):
    st.markdown(f"## {title}")
    simple_explainer("What this section means in simple words", """
This section shows the model through time instead of only today.

You can use it to see whether the signal was strengthening, weakening, or changing direction, and whether gold moved in a similar way.

In plain words: this is the section that helps you understand the **story over time**.
    """)

    # ---- history slicing (same behavior as your previous dropdown)
    res_plot = res_panel.copy()
    try:
        end_dt = res_plot.index.max()
        if history_view == "Last 15y":
            res_plot = res_plot.loc[end_dt - pd.DateOffset(years=15):]
        elif history_view == "Last 5y":
            res_plot = res_plot.loc[end_dt - pd.DateOffset(years=5):]
        elif history_view == "Crisis window only" and win_start and win_end:
            res_plot = res_plot.loc[pd.to_datetime(win_start):pd.to_datetime(win_end)]
        elif history_view == "Crisis window ±5y" and win_start and win_end:
            start_dt = pd.to_datetime(win_start) - pd.DateOffset(years=5)
            end_win = pd.to_datetime(win_end) + pd.DateOffset(years=5)
            res_plot = res_plot.loc[start_dt:end_win]
    except Exception:
        pass

    # ---- Row 1: Signal (left) + Gold vs Signal (right)
    h1, h2 = st.columns(2)

    with h1:
        st.markdown("**Signal over time (with triggers)**")
        fig, ax = plt.subplots(figsize=(7, 3.2))

        ax.plot(res_plot.index, res_plot["SIGNAL"].values, label="Gold Signal", color="tab:blue")

        if res_market is not None and not res_market.empty and "SIGNAL" in res_market.columns:
            market_plot = res_market.reindex(res_plot.index)
            ax.plot(
                market_plot.index,
                market_plot["SIGNAL"],
                label="Market Signal",
                linestyle="--",
                color=theme.market_signal
            )

        ax.axhline(
            trig_hi,
            linestyle="--",
            color=theme.bull_color,
            label=f"Bull trigger ({trig_hi:.2f})"
        )

        ax.axhline(
            trig_lo,
            linestyle="--",
            color=theme.bear_color,
            label=f"Bear trigger ({trig_lo:.2f})"
        )
        ax.axhline(0.0, linewidth=0.8, color="black", label="Zero")

        ax.set_xlabel("Date")
        ax.set_ylabel("Signal (weighted sum)")
        ax.legend(loc="best", fontsize=8)
        ax.grid(True, linewidth=0.4, alpha=0.4)

        plt.tight_layout()
        st.pyplot(fig, clear_figure=True)

    with h2:
        st.markdown("**Gold vs Signal (if available)**")
        gold_col = None
        for c in ["GOLD_USD", "GC=F", "GLD"]:
            if c in res_plot.columns:
                gold_col = c
                break

        if not gold_col:
            st.warning("Gold series not available in this run.")
        else:
            plot_df = res_plot[[gold_col, "SIGNAL"]].dropna().sort_index()
            if plot_df.empty:
                st.warning("Gold vs Signal chart is empty (no overlapping monthly data).")
            else:
                fig, ax1 = plt.subplots(figsize=(7, 3.2))

                # Gold line (gold color)
                l1 = ax1.plot(
                    plot_df.index, plot_df[gold_col].values,
                    linewidth=1.8, label="Gold (USD)", color="gold"
                )

                ax2 = ax1.twinx()
                l2 = ax2.plot(
                    plot_df.index, plot_df["SIGNAL"].values,
                    linewidth=1.2, label="Gold Signal", color="tab:blue"
                )

                l3 = []
                if res_market is not None and not res_market.empty and "SIGNAL" in res_market.columns:
                    market_plot = res_market.reindex(res_plot.index)
                    l3 = ax2.plot(
                        market_plot.index,
                        market_plot["SIGNAL"],
                        label="Market Signal",
                        linestyle="--",
                        color=theme.market_signal
                    )

                lines = l1 + l2 + l3
                labels_ = [ln.get_label() for ln in lines]
                ax1.legend(lines, labels_, loc="best", fontsize=8)

                ax1.grid(True, linewidth=0.4, alpha=0.4)
                plt.tight_layout()
                st.pyplot(fig, clear_figure=True)

    # ---- Row 2: Contributions (left) + Outcome stats (right)
    h3, h4 = st.columns(2)

    with h3:
        st.markdown("**Contributions (stacked by indicator)**")
        contrib_cols = [c for c in res_plot.columns if c.endswith("_CONTRIB")]
        contrib = res_plot[contrib_cols].dropna(how="all")

        if contrib.empty:
            st.info("No contribution history available.")
        else:
            # lookback can be None ("All available")
            if lookback is None:
                contrib = contrib.tail(300)
            else:
                contrib = contrib.tail(max(int(lookback), 300))

            fig, ax = plt.subplots(figsize=(7, 3.2))
            stacks = ax.stackplot(contrib.index, contrib.T.values)

            ax.axhline(0.0, linewidth=0.8, color="black")
            ax.set_xlabel("Date")
            ax.set_ylabel("Contribution to signal")
            ax.grid(True, linewidth=0.4, alpha=0.4)

            # Legend: use column names without "_CONTRIB"
            base_keys = [c.replace("_CONTRIB", "") for c in contrib_cols]
            leg_labels = [labels.get(k, k) for k in base_keys]
            ax.legend(stacks, leg_labels, loc="upper left", fontsize=7, ncols=1)

            plt.tight_layout()
            st.pyplot(fig, clear_figure=True)

    with h4:
        st.markdown("**Outcome stats (light backtest)**")

        # --- Try to use the existing stats logic from your app (table output)
        stats = None

        # 1) If you have a function for it, call it (adjust name if needed)
        if "outcome_stats_table" in globals():
            try:
                stats = outcome_stats_table(res_plot)
            except Exception:
                stats = None
        if stats is None and "compute_outcome_stats" in globals():
            try:
                stats = compute_outcome_stats(res_plot)
            except Exception:
                stats = None

        # 2) Render stats as a themed Plotly table (prevents "white table" on dark themes)
        if isinstance(stats, pd.DataFrame):
            stats_df = stats.copy()
        elif isinstance(stats, dict):
            stats_df = pd.DataFrame([stats])
        elif isinstance(stats, pd.Series):
            stats_df = stats.to_frame().T
        else:
            tmp = res_plot["SIGNAL"].dropna() if "SIGNAL" in res_plot.columns else pd.Series(dtype=float)
            stats_df = pd.DataFrame([{
                "months": int(tmp.shape[0]),
                "mean_signal": float(tmp.mean()) if len(tmp) else np.nan,
                "std_signal": float(tmp.std()) if len(tmp) else np.nan
            }])

        styled_stats = style_signal_table(stats_df, theme)
        st.dataframe(styled_stats, use_container_width=True)

        # st.plotly_chart(
        #     plotly_table(stats_df, theme, "Outcome stats (light backtest)", height=210),
        #     use_container_width=True,
        #     config={"displayModeBar": False},
        #)


def render_gold_action_panels(
        res: pd.DataFrame,
        res_base: pd.DataFrame,
        df: pd.DataFrame,
        mode: str,
        win_start: str | None,
        win_end: str | None,
        labels: dict,
        thresholds: dict,
        core_keys: list,
        latest: pd.Series,
        lookback: int | None,
        show_indicator_thresholds: bool,
        history_view: str,
        trig_hi: float,
        trig_lo: float,
        persist: int,
        latest_market=None,
        core_keys_market=None,
        thresholds_market=None,
        res_market=None,
):
    # -------------------------
    # Indicators (latest)
    # -------------------------
    st.markdown("### Indicators (latest)")

    def badge(state: int) -> str:
        return {+1: "🟢 Bullish", 0: "🟡 Neutral", -1: "🔴 Bearish"}.get(int(state), "—")

    def chunked(lst, n):
        for i in range(0, len(lst), n):
            yield lst[i:i + n]

    update_source = {
        "REAL_YIELD_CPI": "REAL_YIELD_CPI",
        "CPI_YOY": "CPI_YOY",
        "USD_12M_CHG": "USD_TWEX_SPLICE",
        "CURVE_10Y_3M": "CURVE_10Y_3M",
        "DEFICIT_GDP": "DEFICIT_GDP",
        "REAL_YIELD_TIPS10": "REAL_YIELD_TIPS10",
        "HY_OAS": "HY_OAS",
        "QQQ_ABOVE_MA200": "QQQ_ABOVE_MA200",
        "QQQ_MA50_SLOPE_20D": "QQQ_MA50_SLOPE_20D",
        "MARKET_BREADTH_ABOVE_MA200": "MARKET_BREADTH_ABOVE_MA200",
    }

    if mode == "Market Acceleration (fast)":
        update_source.update({
            "GOLD_6M_RET": "GOLD_USD",
            "REALYIELD_3M_CHG": "REAL_YIELD_TIPS10",
            "STRESS_3M_CHG": "HY_OAS",
            "USD_3M_CHG": "USD_TWEX_SPLICE",
        })

    common_macro_keys = ["REAL_YIELD_CPI", "USD_12M_CHG", "CURVE_10Y_3M"]
    gold_only_keys = ["CPI_YOY", "DEFICIT_GDP", "REAL_YIELD_TIPS10"]
    market_only_keys = ["HY_OAS", "QQQ_ABOVE_MA200", "QQQ_MA50_SLOPE_20D", "MARKET_BREADTH_ABOVE_MA200"]

    def render_indicator_group(group_title, cols_to_show):
        shown_cols = [c for c in cols_to_show if c in labels]
        if not shown_cols:
            return

        st.markdown(f"#### {group_title}")

        for row in chunked(shown_cols, 5):
            cards = st.columns(5)
            for i, col in enumerate(row):
                with cards[i]:
                    gold_scored = col in core_keys
                    market_scored = col in core_keys_market if core_keys_market is not None else False
                    is_scored = gold_scored or market_scored

                    gold_state = int(latest.get(col + "_STATE", 0)) if gold_scored else None
                    gold_contrib = float(latest.get(col + "_CONTRIB", 0.0)) if gold_scored else None

                    market_state = (
                        int(latest_market.get(col + "_STATE", 0))
                        if (latest_market is not None and market_scored)
                        else None
                    )

                    market_contrib = (
                        float(latest_market.get(col + "_CONTRIB", 0.0))
                        if (latest_market is not None and market_scored)
                        else None
                    )

                    st.markdown(f"**{labels[col]}**")
                    if not is_scored:
                        st.caption("Context only (not scored in this mode)")

                    val = latest.get(col, np.nan)

                    src_col = update_source.get(col, col)
                    source_df = res_base if src_col in res_base.columns else res

                    if mode == "Market Acceleration (fast)" and src_col == "REAL_YIELD_TIPS10" and src_col not in source_df.columns:
                        if "REAL_YIELD_CPI" in res_base.columns:
                            source_df = res_base
                            src_col = "REAL_YIELD_CPI"

                    if src_col not in source_df.columns:
                        last_val, last_dt = (np.nan, "—")
                    else:
                        last_val, last_dt = latest_value_and_date(source_df, src_col)

                    if pd.isna(val):
                        st.write(f"Latest: **{last_val:.2f}**")
                        st.caption(f"Updated: {last_dt} (stale)")
                    else:
                        if col == "QQQ_ABOVE_MA200":
                            human_val = "Yes" if float(val) >= 0.5 else "No"
                            st.write(f"Latest: **{human_val}**")
                        else:
                            st.write(f"Latest: **{val:.2f}**")
                        st.caption(f"Updated: {last_dt}")

                    # Clarify model usage for HY OAS
                    if col == "HY_OAS":
                        if include_hy:
                            st.caption("Used in: Market model & Gold model (optional — enabled)")
                        else:
                            st.caption("Used in: Market model & Gold model (optional — disabled)")

                    if group_title == "Common Macro Indicators":

                        st.write(
                            f"Gold State: **{badge(gold_state)}**"
                            if gold_state is not None
                            else "Gold State: —"
                        )

                        st.write(
                            f"Market State: **{badge(market_state)}**"
                            if market_state is not None
                            else "Market State: —"
                        )

                        st.write(
                            f"Gold Contrib: **{gold_contrib:+.2f}**"
                            if gold_contrib is not None
                            else "Gold Contrib: —"
                        )

                        st.write(
                            f"Market Contrib: **{market_contrib:+.2f}**"
                            if market_contrib is not None
                            else "Market Contrib: —"
                        )

                    else:
                        single_state = gold_state if gold_state is not None else market_state
                        single_contrib = gold_contrib if gold_contrib is not None else market_contrib

                        if single_state is None:
                            st.write("State: — (not scored in this mode)")
                        else:
                            st.write(f"State: **{badge(single_state)}**")

                        if single_contrib is None:
                            st.write("Contrib: —")
                        else:
                            st.write(f"Contrib: **{single_contrib:+.2f}**")

                    if col in res.columns:
                        s = res[col]
                        spark = s if lookback is None else s.tail(lookback)
                        fig, ax = plt.subplots(figsize=(3.5, 3.2))
                        if col == "QQQ_ABOVE_MA200":
                            ax.step(spark.index, spark.values, where="post", linewidth=1.5, label=labels.get(col, col))
                            ax.set_ylim(-0.1, 1.1)
                            ax.set_yticks([0, 1])
                            ax.set_yticklabels(["Below MA200", "Above MA200"])
                        else:
                            ax.plot(spark.index, spark.values, linewidth=1.2, label=labels.get(col, col))

                        if show_indicator_thresholds and is_scored:
                            if market_scored and not gold_scored:
                                add_indicator_threshold_lines(ax, col, thresholds_market, mode)
                            else:
                                add_indicator_threshold_lines(ax, col, thresholds, mode)

                        leg = ax.legend(loc="best", fontsize=9, frameon=False)

                        # Bloomberg only: make legend text orange
                        theme_label = str(theme) if theme is not None else ""
                        if "Bloomberg" in theme_label:
                            for txt in leg.get_texts():
                                txt.set_color("#FF8C00")
                        ax.grid(True, linewidth=0.4, alpha=0.4)
                        plt.tight_layout()
                        st.pyplot(fig, clear_figure=True)

    render_indicator_group("Common Macro Indicators", common_macro_keys)
    render_indicator_group("Gold-Specific Indicators", gold_only_keys)
    render_indicator_group("Market-Specific Indicators", market_only_keys)

    # -------------------------
    # Triggers (action layer)
    # -------------------------
    st.markdown("### Triggers (action layer)")
    sig = res["SIGNAL"]
    bull_now = sig.iloc[-persist:].ge(trig_hi).all()
    bear_now = sig.iloc[-persist:].le(trig_lo).all()

    tc1, tc2, tc3 = st.columns([1.2, 1.2, 2.2])
    tc1.metric("Bull trigger", "ON ✅" if bull_now else "OFF", f"{persist}m persistence")
    tc2.metric("Bear trigger", "ON ✅" if bear_now else "OFF", f"{persist}m persistence")

    if bull_now:
        tc3.success(f"Action: consider **increasing** gold exposure (Signal ≥ {trig_hi:.2f} for {persist} months).")
    elif bear_now:
        tc3.error(f"Action: consider **reducing** gold exposure (Signal ≤ {trig_lo:.2f} for {persist} months).")
    else:
        tc3.info("Action: no trigger. Maintain baseline / monitor drivers.")

    # -------------------------
    # History
    # -------------------------
    render_history_panel(
        "History",
        res,
        df,
        win_start,
        win_end,
        trig_hi,
        trig_lo,
        res_market=res_market
    )
    return bull_now, bear_now

# ----------------------------
# Structural/Crisis: render gold action panels normally (no tabs)
# ----------------------------
if mode != "Market Acceleration (fast)":
    bull_now, bear_now = render_gold_action_panels(
        res=res,
        res_base=res_base,
        df=df,
        mode=mode,
        win_start=win_start if mode == "Crisis Similarity (template)" else None,
        win_end=win_end if mode == "Crisis Similarity (template)" else None,
        labels=labels,
        thresholds=thresholds,
        core_keys=core_keys,
        latest=latest,
        lookback=lookback,
        show_indicator_thresholds=show_indicator_thresholds,
        history_view=history_view,
        trig_hi=trig_hi,
        trig_lo=trig_lo,
        persist=persist,
        latest_market=latest_market,
        core_keys_market=core_keys_market,
        thresholds_market=thresholds_market,
        res_market=res_market,
    )
# --- Combined charts: (Left) RSI daily+monthly, (Right) Gold price
if mode == "Market Acceleration (fast)":
    tab1, tab2, tab3, tab4 = st.tabs(["Gold Acceleration", "Intraday RSI Screener", "Monte Carlo", "Walk-Forward Backtest"])

    # ----------------------------
    # TAB 1: GOLD ONLY
    # ----------------------------
    with tab1:
        simple_explainer("What this tab means in simple words", """
This tab is the **gold-focused dashboard**.

It combines slow macro signals, acceleration, RSI behavior, and gold price charts so you can judge whether gold looks healthy, stretched, weak, or mixed.

Think of it as the main screen for understanding **gold itself**.
        """)
        left, right = st.columns(2)

        # ---------- Build plot window ----------
        rsi14d_daily = res_base.attrs.get("RSI_14D_DAILY", None)
        gold_daily = res_base.attrs.get("gold_daily", None)

        rsi14m_raw = res_base.attrs.get("RSI_14_ASOF_RAW", None)

        if isinstance(rsi14m_raw, pd.Series) and not rsi14m_raw.dropna().empty:
            rsi14m = rsi14m_raw.dropna().sort_index()
            rsi14m_label = "RSI (14M, as-of today)"
            rsi14m_is_asof = True
        else:
            rsi14m = (res_base["RSI_14"].dropna().sort_index()
                      if "RSI_14" in res_base.columns else pd.Series(dtype=float))
            rsi14m_label = "RSI (14M, month-end)"
            rsi14m_is_asof = False

        rsi14d_full = (rsi14d_daily.dropna().sort_index()
                       if isinstance(rsi14d_daily, pd.Series) else pd.Series(dtype=float))
        gold_d_full = (gold_daily.dropna().sort_index()
                       if isinstance(gold_daily, pd.Series) else pd.Series(dtype=float))
        gold_m_full = (res_base["GOLD_USD"].dropna().sort_index()
                       if "GOLD_USD" in res_base.columns else pd.Series(dtype=float))

        if lookback is None:
            end_dt = gold_d_full.index.max() if not gold_d_full.empty else (
                rsi14d_full.index.max() if not rsi14d_full.empty else rsi14m.index.max()
            )
            start_dt = None
        else:
            end_dt = gold_d_full.index.max() if not gold_d_full.empty else (
                rsi14d_full.index.max() if not rsi14d_full.empty else rsi14m.index.max()
            )

            if isinstance(end_dt, (float, int)):
                # interpret as matplotlib date number (days)
                import matplotlib.dates as mdates

                end_dt = pd.Timestamp(mdates.num2date(end_dt)).tz_localize(None)
            else:
                end_dt = pd.Timestamp(end_dt)
            start_dt = end_dt - pd.DateOffset(months=int(lookback))

        if start_dt is None:
            rsi14d_win = rsi14d_full
            gold_d_win = gold_d_full
            rsi14m_win = rsi14m
            gold_m_win = gold_m_full
        else:
            rsi14d_win = rsi14d_full.loc[rsi14d_full.index >= start_dt] if not rsi14d_full.empty else rsi14d_full
            gold_d_win = gold_d_full.loc[gold_d_full.index >= start_dt] if not gold_d_full.empty else gold_d_full
            rsi14m_win = rsi14m.loc[rsi14m.index >= start_dt] if not rsi14m.empty else rsi14m
            gold_m_win = gold_m_full.loc[gold_m_full.index >= start_dt] if not gold_m_full.empty else gold_m_full

        # Add "current point" to monthly gold so month-end line reaches today
        if not gold_d_win.empty and not gold_m_win.empty:
            gold_m_like = gold_m_win.copy()
            gold_m_like.loc[gold_d_win.index.max()] = float(gold_d_win.iloc[-1])
            gold_m_win = gold_m_like.sort_index()

        # ---------- ROW 1 LEFT: RSI daily + monthly ----------
        with left:
            fig, ax = plt.subplots(figsize=(6.0, 3.2))

            if isinstance(rsi14d_win, pd.Series) and not rsi14d_win.empty:
                ax.plot(rsi14d_win.index, rsi14d_win.values, linewidth=1.5, label="RSI (14D)")

            if isinstance(rsi14m_win, pd.Series) and not rsi14m_win.empty:
                if rsi14m_is_asof:
                    ax.plot(rsi14m_win.index, rsi14m_win.values, linewidth=1.2, linestyle="--", label=rsi14m_label)
                else:
                    ax.plot(rsi14m_win.index, rsi14m_win.values, linewidth=1.2, label=rsi14m_label)

            ax.axhline(70, linestyle="--", linewidth=1.0, alpha=0.7, label="70")
            ax.axhline(30, linestyle="--", linewidth=1.0, alpha=0.7, label="30")
            ax.axhline(85, linestyle=":", linewidth=1.0, alpha=0.6, label="85", color="red")
            ax.axhline(90, linestyle=":", linewidth=1.0, alpha=0.6, label="90", color="red")

            ax.set_title("RSI (Daily + Monthly)", fontsize=11)
            ax.set_ylabel("RSI")
            ax.set_ylim(0, 100)
            ax.grid(True, alpha=0.3)
            ax.legend(loc="best", fontsize=8)

            ax.xaxis.set_major_locator(mdates.AutoDateLocator())
            ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
            fig.autofmt_xdate(rotation=30, ha="right")

            st.pyplot(fig, use_container_width=True)

        # ---------- ROW 1 RIGHT: Gold price ----------
        with right:
            fig, ax = plt.subplots(figsize=(6.0, 3.2))

            if isinstance(gold_d_win, pd.Series) and not gold_d_win.empty:
                ax.plot(gold_d_win.index, gold_d_win.values, linewidth=1.5, label="Gold (daily)")

            if isinstance(gold_m_win, pd.Series) and not gold_m_win.empty:
                ax.plot(gold_m_win.index, gold_m_win.values, marker="o", linewidth=1.2, label="Gold (month-end)")

            ax.set_title("Gold Price (Daily + Month-end)", fontsize=11)
            ax.set_ylabel("USD")
            ax.grid(True, alpha=0.3)
            ax.legend(loc="best", fontsize=8)

            ax.xaxis.set_major_locator(mdates.AutoDateLocator())
            ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
            fig.autofmt_xdate(rotation=30, ha="right")

            st.pyplot(fig, use_container_width=True)

        # ---------- ROW 2: Intraday RSI + Intraday Gold (optional) ----------
        if show_intraday_rsi:
            rsi_intra = res_base.attrs.get("RSI_14_INTRA", None)
            gold_intra = res_base.attrs.get("gold_intraday", None)

            if not (isinstance(rsi_intra, pd.Series) and not rsi_intra.dropna().empty) or not (
                isinstance(gold_intra, pd.Series) and not gold_intra.dropna().empty
            ):
                st.info("Intraday view unavailable (market closed, rate-limited, or no intraday for selected ticker).")
            else:
                rsi_intra = rsi_intra.dropna().sort_index()
                gold_intra = gold_intra.dropna().sort_index()

                end_i = min(rsi_intra.index.max(), gold_intra.index.max())
                start_i = end_i - pd.Timedelta(days=int(intraday_lookback_days))

                rsi_intra_win = rsi_intra.loc[(rsi_intra.index >= start_i) & (rsi_intra.index <= end_i)]
                gold_intra_win = gold_intra.loc[(gold_intra.index >= start_i) & (gold_intra.index <= end_i)]

                i_left, i_right = st.columns(2)

                with i_left:
                    last_rsi = float(rsi_intra_win.iloc[-1]) if not rsi_intra_win.empty else np.nan
                    st.markdown(
                        f"### Intraday RSI (14) - Gold — {intraday_interval}  \nLatest: **{last_rsi:.1f}**"
                        if not np.isnan(last_rsi)
                        else f"### Intraday RSI (14) - Gold — {intraday_interval}"
                    )

                    fig, ax = plt.subplots(figsize=(6.0, 3.2))
                    ax.plot(rsi_intra_win.index, rsi_intra_win.values, linewidth=1.2, label=f"RSI (14, {intraday_interval})")

                    # Monthly RSI as-of-today overlay line
                    rsi_month_asof = res_base.attrs.get("RSI_14_ASOF_RAW", None)
                    if isinstance(rsi_month_asof, pd.Series) and not rsi_month_asof.dropna().empty:
                        rsi_month_asof_val = float(rsi_month_asof.dropna().iloc[-1])
                        ax.axhline(rsi_month_asof_val, linestyle="--", linewidth=1.4, alpha=0.85, color="tab:orange",
                                   label=f"Monthly RSI as-of-today ({rsi_month_asof_val:.1f})")

                    ax.axhline(70, linestyle="--", linewidth=1.0, alpha=0.7, label="70")
                    ax.axhline(30, linestyle="--", linewidth=1.0, alpha=0.7, label="30")
                    ax.axhline(85, linestyle=":", linewidth=1.0, alpha=0.6, label="85", color="red")
                    ax.axhline(90, linestyle=":", linewidth=1.0, alpha=0.6, label="90", color="red")

                    ax.set_ylabel("RSI")
                    ax.set_ylim(0, 100)
                    ax.grid(True, alpha=0.3)
                    ax.legend(loc="best", fontsize=8)

                    ax.xaxis.set_major_locator(mdates.AutoDateLocator())
                    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d %H:%M"))
                    fig.autofmt_xdate(rotation=30, ha="right")

                    st.pyplot(fig, use_container_width=True)

                with i_right:
                    last_px = float(gold_intra_win.iloc[-1]) if not gold_intra_win.empty else np.nan
                    st.markdown(
                        f"### Intraday - Gold — {intraday_interval}  \nLatest: **{last_px:,.2f}**"
                        if not np.isnan(last_px)
                        else f"### Intraday - Gold — {intraday_interval}"
                    )

                    fig, ax = plt.subplots(figsize=(6.0, 3.2))
                    ax.plot(gold_intra_win.index, gold_intra_win.values, linewidth=1.2, label=f"Gold ({intraday_interval})")

                    ax.set_ylabel("USD")
                    ax.grid(True, alpha=0.3)
                    ax.legend(loc="best", fontsize=8)

                    ax.xaxis.set_major_locator(mdates.AutoDateLocator())
                    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d %H:%M"))
                    fig.autofmt_xdate(rotation=30, ha="right")

                    st.pyplot(fig, use_container_width=True)

    # ----------------------------
    # TAB 2: TICKERS ONLY (NO GOLD)
    # ----------------------------
    with tab2:

        with st.expander("### How to interpret this tab"):
            st.markdown("""
    This section answers the question:
    **Is there a short-term trading setup forming?**

    It analyzes short-term momentum signals using RSI and volatility regimes.

    Interpretation guide:

    • **BUY / SELL** → actionable short-term trade signal  
    • **WATCH** → setup forming but not yet confirmed  
    • **NEUTRAL** → no meaningful tactical signal  

    This tab focuses on **entry timing**, not medium-term position decisions.
    """)

        intraday_screener.render_intraday_rsi_screener_tab(
            rsi_func=rsi,
            zscore_func=zscore,
            tickers=tickers_top10,
            refresh_token=refresh_token
        )
    # ----------------------------
    # TAB 3: Monte Carlo
    # ----------------------------
    with tab3:

        if mode == "Crisis Similarity (template)":
            st.info("Monte Carlo disabled in Crisis Similarity mode.")
        else:
            st.subheader("Monte Carlo — Analysis")
            simple_explainer("What this section means in simple words", """
This section explores **what could happen next**, not just what happened so far.

It creates many possible future price paths based on similar past market states. The goal is to show a reasonable range of outcomes, including a typical path, a weaker path, and a stronger path.
            """)

            with st.expander("### How to interpret these charts"):
                st.markdown("""

              This Section is answering the question:
              **Should I open or hold a position over the next weeks or months?**
              
              The Monte Carlo simulation estimates possible price paths based on the
              current tactical regime.              
              This chart shows how the ticker's price behaved in the past when the economy looked like it does today.  
               - The middle line (p50) shows what usually happened.
               - The lower line  (p10) shows a realistic downside scenario.
               - The upper line  (p90) shows a realistic upside scenario.
              A wider band means more uncertainty. A narrower band means more stability.

              **1) General Direction** — Is the ticker's price usually rising or falling in this environment?
              p50 curve slopes upward: In similar economic conditions in the past, gold usually moved higher.  
              p50 slopes downward: In similar conditions, gold usually moved lower.  
              **2) Risk / volatility regime**  
              If the gap between p10 and p90 is wide, this environment historically had high volatility / uncertainty.  
              If it’s narrow, this environment hws usually been calmer and more predictable.  
              **3) Downside risk - What could go wrong?**  
              Look at where p10 is after 6–12 months.  
              A realistic bad outcome(worst case, not collapse not crisis), but a meaningful drop that has happened before in similar conditions.
              **4) Upside potential** - What could go well?  
              Look at p90 after 6–12 months.  
              A realistic strong outcome — not a bubble or mania, but a good rally that has happened before.  

              **Simulation type:** Tactical per-ticker Monte Carlo  
              **Horizon:** selected number of bars  
              **Simulations:** user-selected  
              """)

            # --- selection UI
            left, right = st.columns([1, 3], gap="large")
            with left:
                overlay = st.checkbox("Overlay view (compare selected tickers)", value=True)
                mc_steps = st.slider("Simulation horizon (bars)", 20, 250, 60, 10)
                mc_n2 = st.slider("Simulations", 500, 5000, 2000, 500)

                # Use your existing ticker universe
                universe = list(dict.fromkeys((tickers_top10 or []) + ["GLD", "SPY", "QQQ"]))
                universe = [t for t in universe if isinstance(t, str) and t.strip()]

                if "mc_selected" not in st.session_state:
                    st.session_state["mc_selected"] = ["GLD", "SPY"]

                selected = []
                for t in universe:
                    if st.checkbox(t, value=(t in st.session_state["mc_selected"]), key=f"mc_cb_{t}"):
                        selected.append(t)
                st.session_state["mc_selected"] = selected

            with right:
                if not selected:
                    st.info("Select tickers on the left.")
                    st.stop()

                # --- Fetch DAILY closes (2y) + build RSI + slope-based tactical state
                # Reuse intraday_screener's yfinance helper for DAILY prices
                close_d = intraday_screener.yf_multi_close_fixed_period(
                    selected, interval="1d", period="2y", refresh_token=refresh_token
                )

                if close_d is None or close_d.empty:
                    st.warning("No price data available for selected tickers.")
                    st.stop()

                p50_overlay = {}

                # ---- Fetch QQQ once (used for RS vs QQQ) ----
                try:
                    qqq_df = intraday_screener.yf_multi_close_fixed_period(
                        ["QQQ"],
                        interval="1d",
                        period="2y",
                        refresh_token=refresh_token
                    )
                    qqq_px = qqq_df["QQQ"].dropna() if (
                                qqq_df is not None and not qqq_df.empty and "QQQ" in qqq_df.columns) else pd.Series(
                        dtype=float)
                except Exception:
                    qqq_px = pd.Series(dtype=float)

                mc_summary_rows = []

                for t in selected:
                    px = close_d[t].dropna() if t in close_d.columns else pd.Series(dtype=float)
                    if px.empty or px.shape[0] < 60:
                        st.warning(f"{t}: not enough daily history for RSI/state.")
                        continue

                    r = rsi(px, period=14).dropna()
                    r = r.reindex(px.index, method="ffill")
                    slope = r.diff(5)  # 5 trading days slope
                    # --- Trend filters ---
                    ma200 = px.rolling(200).mean()
                    above_ma200 = px > ma200

                    ma50 = px.rolling(50).mean()
                    ma50_slope = ma50.pct_change(20)

                    # --- Volatility regime ---
                    vol = px.pct_change().rolling(20).std() * np.sqrt(252)
                    vol_q75 = vol.dropna().quantile(0.75) if not vol.dropna().empty else np.nan
                    vol_regime = pd.Series(
                        np.where(vol > vol_q75, "HIGH", "LOW"),
                        index=px.index
                    )

                    # --- Relative strength vs QQQ ---
                    if not qqq_px.empty:
                        rs_vs_qqq = px.pct_change(63) - qqq_px.pct_change(63)
                        rs_vs_qqq = rs_vs_qqq.reindex(px.index)
                    else:
                        rs_vs_qqq = pd.Series(index=px.index, dtype=float)

                    state = pd.Series(
                        [
                            mc.classify_accel_state(
                                rv,
                                sv,
                                above_ma200.iloc[i] if i < len(above_ma200) else None,
                                rs_vs_qqq.iloc[i] if i < len(rs_vs_qqq) else None,
                                ma50_slope.iloc[i] if i < len(ma50_slope) else None,
                                vol_regime.iloc[i] if i < len(vol_regime) else None,
                            )
                            for i, (rv, sv) in enumerate(zip(r.values, slope.values))
                        ],
                        index=px.index,
                        name="STATE",
                    )

                    paths, state_now = mc.monte_carlo_paths_by_tactical_state_block(
                        price=px,
                        state=state,
                        horizon_steps=mc_steps,
                        n_sims=mc_n2,
                        block_size=5,
                        seed=7,
                    )

                    if paths is None:
                        st.warning(f"{t}: not enough history in current tactical state ({state_now}).")
                        continue

                    bands = mc.mc_percentiles(paths)
                    # store p50 for overlay
                    p50_overlay[t] = bands["p50"].rename(t)

                    if not overlay:
                        st.markdown(f"### {t} (state: `{state_now}`)")
                        st.line_chart(bands)

                    # --- Horizon summary stats (always compute for combined table) ---
                    start_price = float(px.iloc[-1])

                    p10_end = float(bands["p10"].iloc[-1])
                    p50_end = float(bands["p50"].iloc[-1])
                    p90_end = float(bands["p90"].iloc[-1])

                    p10_ret = (p10_end / start_price - 1.0) * 100.0
                    p50_ret = (p50_end / start_price - 1.0) * 100.0
                    p90_ret = (p90_end / start_price - 1.0) * 100.0

                    prob_up = float((paths[:, -1] > start_price).mean()) * 100.0

                    if t.upper() == "GLD":
                        structural_regime_for_ticker = str(latest.get("REGIME", "—"))  # gold structural regime
                    else:
                        structural_regime_for_ticker = market_regime_now  # market structural regime

                    decision = decide_trade(
                        ticker=t,
                        structural_regime=structural_regime_for_ticker,
                        tactical_state=state_now,
                        mc_typical_pct=p50_ret,
                        mc_prob_higher_pct=prob_up,
                        above_ma200=bool(above_ma200.iloc[-1]) if len(above_ma200.dropna()) else None,
                    )

                    mc_summary_rows.append({
                        "Ticker": t,
                        "Structural Regime": structural_regime_for_ticker,
                        "State": state_now,
                        "Decision": decision.action,
                        "Confidence (%)": round(decision.confidence * 100, 1),
                        "Position Size (%)": round(decision.position_size_pct * 100, 1),
                        "Current Price": round(start_price, 2),
                        f"Downside (p10, {mc_steps} bars)": round(p10_end, 2),
                        f"Typical (p50, {mc_steps} bars)": round(p50_end, 2),
                        f"Upside (p90, {mc_steps} bars)": round(p90_end, 2),
                        "Downside %": round(p10_ret, 2),
                        "Typical %": round(p50_ret, 2),
                        "Upside %": round(p90_ret, 2),
                        "Prob. Finish Higher (%)": round(prob_up, 1),
                        "Reason": decision.reason,
                    })

                if overlay:
                    if not p50_overlay:
                        st.warning("Nothing to overlay (no ticker had enough state history).")
                    else:
                        df_overlay = pd.concat(p50_overlay.values(), axis=1)
                        st.markdown("### Overlay (Median path only)")
                        st.caption(
                            "Each line is the typical (p50) simulated path under each ticker’s current tactical state.")
                        st.line_chart(df_overlay)

                if mc_summary_rows:
                    st.markdown("### Monte Carlo Summary Table")
                    simple_explainer("What this table means in simple words", """
This table puts the main Monte Carlo results for several tickers in one place.

It helps you compare which assets look more favorable, less favorable, or more uncertain under their current tactical setup.
                    """)

                    mc_summary_df = pd.DataFrame(mc_summary_rows)

                    # Sort by probability of finishing higher
                    mc_summary_df = mc_summary_df.sort_values(
                        "Prob. Finish Higher (%)",
                        ascending=False
                    )

                    best = mc_summary_df.iloc[0]["Ticker"]
                    st.caption(f"Top Monte Carlo probability: **{best}**")
                    st.dataframe(mc_summary_df,use_container_width=True,hide_index=True)

    # ----------------------------
    # TAB 4: Walk-Forward Backtest
    # ----------------------------

    def _prepare_bt_for_portfolio(bt_in: pd.DataFrame, trading_mode: str) -> pd.DataFrame:
        bt = bt_in.copy()

        if "signal_action" not in bt.columns and "decision" in bt.columns:
            bt["signal_action"] = bt["decision"]
        if "signal_confidence" not in bt.columns and "decision_confidence" in bt.columns:
            bt["signal_confidence"] = bt["decision_confidence"]

        bt["signal_action"] = (
            bt["signal_action"]
            .astype(str)
            .str.upper()
            .str.strip()
        )

        # Preserve replay-generated sizing whenever present
        if "position_size_pct" not in bt.columns:
            bt["position_size_pct"] = np.nan

        bt["position_size_pct"] = pd.to_numeric(bt["position_size_pct"], errors="coerce")

        mode = str(trading_mode).lower().strip()
        sell_mask = bt["signal_action"] == "SELL"

        if mode == "long_short":
            # Only fill missing SELL size; do not overwrite existing replay sizing
            bt.loc[sell_mask & bt["position_size_pct"].isna(), "position_size_pct"] = 100.0
        else:
            bt.loc[sell_mask & bt["position_size_pct"].isna(), "position_size_pct"] = 0.0

        # Conservative defaults only when size is missing
        bt.loc[(bt["signal_action"] == "HOLD") & bt["position_size_pct"].isna(), "position_size_pct"] = 0.0
        bt.loc[(bt["signal_action"] == "WATCH") & bt["position_size_pct"].isna(), "position_size_pct"] = 0.0
        bt.loc[(bt["signal_action"] == "BUY") & bt["position_size_pct"].isna(), "position_size_pct"] = 100.0

        return bt

    with tab4:
        st.subheader("Walk-Forward Macro Backtest")
        simple_explainer("What this section means in simple words", """
    This section is a **historical replay**. It takes the model back in time and asks:
    "If I were living at that date, and only knew what was available then, what would the platform have said?"

    So instead of only looking at today, you can see whether the logic was usually helpful, mixed, or weak across many different periods.

    It is still a **research tool**, not a fully realistic trading account simulation.
    It uses simple assumptions for costs, slippage, confirmation, and holding rules so you can compare ideas more fairly.
        """)

        # ----------------------------
        # Core replay settings
        # ----------------------------
        b1, b2, b3, b4, b5 = st.columns(5)
        with b1:
            min_bt_date = pd.to_datetime(res_base.index.min()).date() if len(res_base.index) else datetime(1975, 1,
                                                                                                           1).date()
            max_bt_date = pd.to_datetime(res_base.index.max()).date() if len(
                res_base.index) else datetime.today().date()
            default_bt_date = max(min_bt_date, datetime(2020, 1, 1).date())
            bt_start_dt = st.date_input(
                "Start date",
                value=default_bt_date,
                min_value=min_bt_date,
                max_value=max_bt_date,
                key="bt_start_dt",
            )
            bt_start = pd.Timestamp(bt_start_dt).date().isoformat()

        with b2:
            bt_step = int(st.slider("Replay step (months)", 1, 12, 1, 1, key="bt_step"))

        with b3:
            bt_fwd = int(st.slider("Forward horizon (months)", 1, 12, 12, 1, key="bt_fwd"))

        with b4:
            bt_price_col = st.selectbox(
                "Price series",
                [c for c in ["GOLD_USD", "GLD"] if c in res_base.columns] or ["GOLD_USD"],
                index=0,
                key="bt_price_col",
            )

        with b5:
            bt_initial_capital = float(
                st.number_input(
                    "Initial investment",
                    min_value=100.0,
                    max_value=100000000.0,
                    value=10000.0,
                    step=100.0,
                    key="bt_initial_capital",
                )
            )

        simple_explainer("What forward horizon means", """
    The **x-axis shows decision dates**, not the end date of each trade.

    A **12-month forward horizon** means each replay date is judged by what happened over the **next 12 months**.
    That forward result is used for evaluation tables.

    The **benchmark curve** is different: it is a simple normalized buy-and-hold path of the selected asset from the first replay date onward.
        """)

        c1, c2, c3, c4 = st.columns(4)
        with c4:
            annualize_by = st.selectbox("Analytics basis", ["Monthly", "Custom"], index=0, key="bt_annualize_by")
            periods_per_year = 12.0 if annualize_by == "Monthly" else max(1.0, 12.0 / max(float(bt_step), 1.0))

        # ----------------------------
        # Strategy rules
        # ----------------------------
        st.markdown("### Strategy Rules")
        simple_explainer("What these strategy rules mean", """
    These options let you compare **different decision logics**, not just different trading assumptions.

    - **Current fused strategy** uses the combined macro state.
    - **Structural-only** listens only to the gold structural regime.
    - **Acceleration-only** listens only to the acceleration regime.
    - **RSI-based gold rule** uses gold RSI as a simple timing rule.
    - **Always-long benchmark** stays invested all the time.
        """)

        rule_options = [
            "Current fused strategy",
            "Structural-only",
            "Acceleration-only",
            "RSI-based gold rule",
            "Always-long benchmark",
        ]

        sr1, sr2, sr3 = st.columns([1.2, 1.2, 1.0])
        with sr1:
            bt_primary_rule = st.selectbox("Primary strategy rule", rule_options, index=0, key="bt_primary_rule")
        with sr2:
            bt_compare_rule = st.selectbox("Comparison strategy rule", rule_options, index=3, key="bt_compare_rule")
        with sr3:
            compare_enabled = st.checkbox("Enable side-by-side rule comparison", value=True, key="bt_compare_enabled")

        st.markdown("#### Primary execution settings")
        p1, p2, p3 = st.columns(3)
        with p1:
            bt_trading_mode = st.selectbox("Primary trading mode", ["long_short", "long_flat"], index=0,
                                           key="bt_trading_mode")
        with p2:
            bt_tc_bps = float(
                st.number_input("Primary transaction cost (bps)", min_value=0.0, max_value=500.0, value=10.0, step=1.0,
                                key="bt_tc_bps"))
        with p3:
            bt_slip_bps = float(
                st.number_input("Primary slippage (bps)", min_value=0.0, max_value=500.0, value=5.0, step=1.0,
                                key="bt_slip_bps"))

        if compare_enabled:
            simple_explainer("What the comparison does", """
    This compares either **different rules**, **different trading assumptions**, or both.

    For example, you can compare:
    - fused strategy versus structural-only
    - RSI rule versus always-long
    - long/flat versus long/short under the same rule
            """)

            st.markdown("#### Comparison execution settings")
            cp1, cp2, cp3 = st.columns(3)
            with cp1:
                compare_trading_mode = st.selectbox(
                    "Comparison trading mode",
                    ["long_flat", "long_short"],
                    index=0 if bt_trading_mode == "long_flat" else 1,
                    key="compare_trading_mode",
                )
            with cp2:
                compare_tc_bps = float(st.number_input(
                    "Comparison transaction cost (bps)",
                    min_value=0.0,
                    max_value=500.0,
                    value=max(bt_tc_bps, 10.0),
                    step=1.0,
                    key="compare_tc_bps",
                ))
            with cp3:
                compare_slip_bps = float(st.number_input(
                    "Comparison slippage (bps)",
                    min_value=0.0,
                    max_value=500.0,
                    value=max(bt_slip_bps, 5.0),
                    step=1.0,
                    key="compare_slip_bps",
                ))
        else:
            compare_trading_mode = bt_trading_mode
            compare_tc_bps = bt_tc_bps
            compare_slip_bps = bt_slip_bps

        # ----------------------------
        # Execution realism
        # ----------------------------
        st.markdown("### Execution Realism")
        with st.expander("What these realism settings mean"):
            st.markdown("""
    These settings make the replay more realistic.

    - **Execution lag** delays trades instead of acting instantly.
    - **Confirmation steps** require the same signal to appear repeatedly before acting.
    - **Minimum hold** prevents the strategy from flipping too quickly.
    - **Rebalance mode** controls whether size changes are applied often or only when direction truly changes.
    - **Stop-loss / Take-profit** apply a simple end-of-step risk rule to cap gains or losses for that replay interval.
            """)

        r1, r2, r3 = st.columns(3)
        bt_execution_lag = int(
            r1.number_input("Execution lag (steps)", min_value=0, max_value=6, value=0, step=1, key="bt_execution_lag"))
        bt_confirmation_steps = int(r2.number_input("Confirmation steps", min_value=1, max_value=6, value=1, step=1,
                                                    key="bt_confirmation_steps"))
        bt_min_hold_steps = int(r3.number_input("Minimum hold (steps)", min_value=0, max_value=12, value=0, step=1,
                                                key="bt_min_hold_steps"))

        r4, r5, r6 = st.columns(3)
        bt_rebalance_mode = r4.selectbox("Rebalance mode", ["on_change", "every_signal", "threshold"], index=0,
                                         key="bt_rebalance_mode")
        bt_rebalance_threshold = float(
            r5.number_input("Rebalance threshold (%)", min_value=0.0, max_value=50.0, value=2.5, step=0.5,
                            key="bt_rebalance_threshold"))
        bt_stop_loss_pct = float(r6.number_input("Stop-loss (%)", min_value=0.0, max_value=50.0, value=5.0, step=0.5,
                                                 key="bt_stop_loss_pct"))

        r7, _, _ = st.columns(3)
        bt_take_profit_pct = float(
            r7.number_input("Take-profit (%)", min_value=0.0, max_value=100.0, value=5.0, step=0.5,
                            key="bt_take_profit_pct"))

        # ----------------------------
        # Historical Monte Carlo Validation
        # ----------------------------
        st.markdown("### Historical Monte Carlo Validation")
        simple_explainer("What this section means in simple words", """
        This section checks whether the Monte Carlo forecast method would have been useful in the past.

        For each replay date, the platform:
        - runs a Monte Carlo forecast using only information available then
        - records p10 / p50 / p90
        - compares that forecast with what actually happened later

        So this is not today's forecast. It is a historical test of how well the forecast cone worked.
        """)

        mv1, mv2, mv3, mv4 = st.columns(4)
        with mv1:
            wf_mc_enabled = st.checkbox("Enable MC validation", value=True, key="wf_mc_enabled")
        with mv2:
            wf_mc_hist_bars = int(
                st.number_input(
                    "MC history bars",
                    min_value=120,
                    max_value=2000,
                    value=500,
                    step=20,
                    key="wf_mc_hist_bars",
                )
            )
        with mv3:
            wf_mc_block = int(
                st.number_input(
                    "MC block size",
                    min_value=2,
                    max_value=20,
                    value=5,
                    step=1,
                    key="wf_mc_block",
                )
            )
        with mv4:
            wf_mc_sims = int(
                st.number_input(
                    "MC simulations",
                    min_value=200,
                    max_value=5000,
                    value=1000,
                    step=100,
                    key="wf_mc_sims",
                )
            )

        # ----------------------------
        # Rule decision helpers
        # ----------------------------
        def _decision_from_rule(rule_name, ms, lg, lm, la, hist_base):
            rule_name = str(rule_name)

            if rule_name == "Current fused strategy":
                if ms is None:
                    return {
                        "action": "HOLD",
                        "confidence": 0.20,
                        "position_size_pct": 0.0,
                        "reason": "Macro state unavailable in replay row",
                    }

                dec = decide_trade(ticker=bt_price_col, macro_state=ms)
                pos_size = float(dec.position_size_pct)
                if pos_size <= 1.0:
                    pos_size *= 100.0
                return {
                    "action": dec.action,
                    "confidence": dec.confidence,
                    "position_size_pct": pos_size,
                    "reason": dec.reason,
                }

            if rule_name == "Structural-only":
                sig = pd.to_numeric(pd.Series([lg.get("SIGNAL")]), errors="coerce").iloc[0]
                reg = str(lg.get("REGIME") or "")
                if pd.notna(sig) and float(sig) >= 0.60:
                    return {"action": "BUY", "confidence": 0.75, "position_size_pct": 100.0,
                            "reason": f"Structural regime bullish ({reg})"}
                if pd.notna(sig) and float(sig) >= 0.20:
                    return {"action": "WATCH", "confidence": 0.60, "position_size_pct": 50.0,
                            "reason": f"Structural regime constructive ({reg})"}
                if pd.notna(sig) and float(sig) <= -0.60:
                    return {"action": "SELL", "confidence": 0.75, "position_size_pct": 0.0,
                            "reason": f"Structural regime defensive ({reg})"}
                return {"action": "HOLD", "confidence": 0.45, "position_size_pct": 0.0,
                        "reason": f"Structural regime mixed ({reg})"}

            if rule_name == "Acceleration-only":
                sig = pd.to_numeric(pd.Series([la.get("SIGNAL")]), errors="coerce").iloc[0]
                reg = str(la.get("REGIME") or "")
                if pd.notna(sig) and float(sig) >= 0.60:
                    return {"action": "BUY", "confidence": 0.72, "position_size_pct": 100.0,
                            "reason": f"Acceleration bullish ({reg})"}
                if pd.notna(sig) and float(sig) >= 0.20:
                    return {"action": "WATCH", "confidence": 0.58, "position_size_pct": 50.0,
                            "reason": f"Acceleration constructive ({reg})"}
                if pd.notna(sig) and float(sig) <= -0.60:
                    return {"action": "SELL", "confidence": 0.72, "position_size_pct": 0.0,
                            "reason": f"Acceleration headwind ({reg})"}
                return {"action": "HOLD", "confidence": 0.45, "position_size_pct": 0.0,
                        "reason": f"Acceleration mixed ({reg})"}

            if rule_name == "RSI-based gold rule":
                rsi_val = np.nan
                rsi_slope = np.nan

                for c in ["RSI_14_ASOF", "RSI_14", "RSI_14D"]:
                    if c in hist_base.columns and hist_base[c].dropna().shape[0]:
                        rsi_val = float(hist_base[c].dropna().iloc[-1])
                        break

                if "RSI_SLOPE_3M" in hist_base.columns and hist_base["RSI_SLOPE_3M"].dropna().shape[0]:
                    rsi_slope = float(hist_base["RSI_SLOPE_3M"].dropna().iloc[-1])

                if pd.notna(rsi_val) and rsi_val <= 40:
                    return {"action": "BUY", "confidence": 0.55, "position_size_pct": 100.0,
                            "reason": f"RSI constructive / oversold ({rsi_val:.1f})"}
                if pd.notna(rsi_val) and rsi_val >= 70:
                    return {"action": "SELL", "confidence": 0.55, "position_size_pct": 0.0,
                            "reason": f"RSI stretched / overbought ({rsi_val:.1f})"}
                if pd.notna(rsi_val) and pd.notna(rsi_slope) and rsi_val < 50 and rsi_slope > 0:
                    return {"action": "WATCH", "confidence": 0.50, "position_size_pct": 50.0,
                            "reason": f"RSI improving ({rsi_val:.1f}, slope {rsi_slope:.2f})"}
                return {"action": "HOLD", "confidence": 0.40, "position_size_pct": 0.0, "reason": "RSI neutral / mixed"}

            if rule_name == "Always-long benchmark":
                return {"action": "BUY", "confidence": 1.00, "position_size_pct": 100.0,
                        "reason": "Always invested benchmark"}

            return {"action": "HOLD", "confidence": 0.40, "position_size_pct": 0.0, "reason": "Fallback"}


        def _compute_bt_row_factory(rule_name):
            def _compute_bt_row(hist_df: pd.DataFrame, dt: pd.Timestamp) -> dict:
                if hist_df.empty or len(hist_df.dropna(how="all")) < 36:
                    return {"error": "Not enough history"}

                hist_accel = compute_accel_features(hist_df.copy())

                g_res, _, _, _, _ = run_structural(hist_df.copy(), weights_struct, dirs_struct)
                m_res, _, _, _, _ = run_market_structural(hist_df.copy())
                a_res, _, _, _, _ = run_accel(hist_accel.copy(), accel_method, w_g, w_ry, w_st, w_u)

                lg = g_res.iloc[-1] if not g_res.empty else pd.Series(dtype=float)
                lm = m_res.iloc[-1] if not m_res.empty else pd.Series(dtype=float)
                la = a_res.iloc[-1] if not a_res.empty else pd.Series(dtype=float)

                macro_state_label = "—"
                macro_score_val = np.nan
                ms = None
                macro_error = None

                # Only fused strategy really needs macro_state, but we still try to build it safely
                try:
                    ms = build_macro_state(
                        asset=bt_price_col,
                        as_of_date=str(pd.Timestamp(dt).date()),
                        gold_signal=lg.get("SIGNAL"),
                        gold_regime=lg.get("REGIME"),
                        market_signal=lm.get("SIGNAL"),
                        market_regime=lm.get("REGIME"),
                        accel_signal=la.get("SIGNAL"),
                        accel_regime=la.get("REGIME"),
                    )
                    macro_state_label = getattr(ms, "fused_state", "—")
                    macro_score_val = getattr(ms, "fused_score", np.nan)
                except Exception as e:
                    macro_error = str(e)
                    ms = None

                dec = _decision_from_rule(rule_name, ms, lg, lm, la, hist_df)

                reason_text = dec["reason"]
                if macro_error is not None:
                    reason_text = f"{reason_text} | macro_state_error: {macro_error}"

                return {
                    "gold_signal": lg.get("SIGNAL"),
                    "gold_regime": lg.get("REGIME"),
                    "market_signal": lm.get("SIGNAL"),
                    "market_regime": lm.get("REGIME"),
                    "accel_signal": la.get("SIGNAL"),
                    "accel_regime": la.get("REGIME"),
                    "macro_score": macro_score_val,
                    "macro_state": macro_state_label,
                    "decision": dec["action"],
                    "decision_confidence": dec["confidence"],
                    "position_size_pct": dec["position_size_pct"],
                    "reason": reason_text,
                    "strategy_rule": rule_name,
                }

            return _compute_bt_row


        # ----------------------------
        # Run primary backtest
        # ----------------------------
        bt_cfg = MacroBacktestConfig(
            ticker=bt_price_col,
            start_date=bt_start,
            step=bt_step,
            price_col=bt_price_col,
            forward_return_months=bt_fwd,
            trading_mode=bt_trading_mode,
            transaction_cost_bps=bt_tc_bps,
            slippage_bps=bt_slip_bps,
        )

        bt_df = run_macro_strategy_backtest(
            res_base,
            compute_one_date=_compute_bt_row_factory(bt_primary_rule),
            config=bt_cfg,
        )

        with st.expander("Primary run window", expanded=True):
            if not bt_df.empty and "as_of_date" in bt_df.columns:
                d = pd.to_datetime(bt_df["as_of_date"], errors="coerce")
                st.write({
                    "rows": int(len(bt_df)),
                    "min_date": str(d.min()),
                    "max_date": str(d.max()),
                })

        with st.expander("Replay Decision Debug", expanded=False):

            st.write("Primary bt_df columns:")
            st.write(list(bt_df.columns))

            if "decision" in bt_df.columns:
                st.write("Primary decision counts:")
                st.write(bt_df["decision"].astype(str).value_counts(dropna=False))
                st.write("Primary decision null count:", int(bt_df["decision"].isna().sum()))
            else:
                st.error("Primary bt_df is missing the 'decision' column.")

            debug_cols = [c for c in [
                "as_of_date",
                "strategy_rule",
                "gold_signal",
                "gold_regime",
                "market_signal",
                "market_regime",
                "accel_signal",
                "accel_regime",
                "macro_state",
                "macro_score",
                "decision",
                "decision_confidence",
                "position_size_pct",
                "signal_action",
                "signal_confidence",
                "reason",
                "error",
            ] if c in bt_df.columns]

            if debug_cols:
                st.write("Primary replay sample:")
                st.dataframe(bt_df[debug_cols].head(25), width="stretch", hide_index=True)

            primary_macro_debug_cols = [c for c in [
                "as_of_date",
                "gold_signal",
                "gold_regime",
                "market_signal",
                "market_regime",
                "accel_signal",
                "accel_regime",
                "macro_score",
                "macro_state",
                "decision",
                "reason",
            ] if c in bt_df.columns]

            st.write("Primary macro input sample:")
            st.dataframe(
                bt_df[primary_macro_debug_cols].head(25),
                width="stretch",
                hide_index=True,
            )

            if "error" in bt_df.columns and bt_df["error"].notna().any():
                st.warning("Replay row errors detected in primary strategy")
                st.dataframe(
                    bt_df.loc[bt_df["error"].notna(), ["as_of_date", "error"]].head(20),
                    width="stretch",
                    hide_index=True,
                )

        with st.expander("Primary decision counts", expanded=True):
            if "decision" in bt_df.columns:
                st.write(bt_df["decision"].astype(str).value_counts(dropna=False))

        with st.expander("Primary BUY rows", expanded=False):
            if "decision" in bt_df.columns:
                buy_rows = bt_df[bt_df["decision"].astype(str).str.upper() == "BUY"].copy()

                cols = [c for c in [
                    "as_of_date",
                    "gold_signal",
                    "market_signal",
                    "accel_signal",
                    "macro_score",
                    "macro_state",
                    "decision",
                    "decision_confidence",
                    "reason"
                ] if c in buy_rows.columns]

                st.dataframe(buy_rows[cols], use_container_width=True, hide_index=True)

        with st.expander("Primary WATCH rows", expanded=False):
            if "decision" in bt_df.columns:
                watch_rows = bt_df[bt_df["decision"].astype(str).str.upper() == "WATCH"].copy()

                cols = [c for c in [
                    "as_of_date",
                    "gold_signal",
                    "market_signal",
                    "accel_signal",
                    "macro_score",
                    "macro_state",
                    "decision",
                    "decision_confidence",
                    "reason"
                ] if c in watch_rows.columns]

                st.dataframe(watch_rows[cols], use_container_width=True, hide_index=True)

        bt_df["signal_action"] = bt_df["decision"] if "decision" in bt_df.columns else np.nan
        bt_df["signal_confidence"] = bt_df["decision_confidence"] if "decision_confidence" in bt_df.columns else np.nan

        # ----------------------------
        # Replay dataframe debug
        # ----------------------------
        debug_cols = [c for c in ["as_of_date", "decision", "position_size_pct", "reason", "error"] if
                      c in bt_df.columns]


        mc_validation_df = pd.DataFrame()
        mc_validation_summary = {}

        if wf_mc_enabled:
            mc_cfg = MCValidationConfig(
                horizon_steps=int(bt_fwd),
                n_sims=int(wf_mc_sims),
                block_size=int(wf_mc_block),
                seed=7,
            )

            mc_validation_df = run_walkforward_mc_validation(
                replay_df=bt_df,
                full_price=res_base[bt_price_col],
                benchmark_price=res_base["QQQ"] if "QQQ" in res_base.columns else None,
                horizon_months=int(bt_fwd),
                cfg=mc_cfg,
            )

            if mc_validation_df is not None and not mc_validation_df.empty:
                bt_df = mc_validation_df
                mc_validation_summary = summarize_mc_validation(mc_validation_df)

        with st.expander("Primary pre-portfolio debug", expanded=False):
            st.write("Primary bt_df columns before _prepare_bt_for_portfolio:")
            st.write(list(bt_df.columns))

        bt_df_for_port = _prepare_bt_for_portfolio(bt_df, bt_trading_mode)

        with st.expander("DEBUG SUMMARY BEFORE PORTFOLIO (new)", expanded=True):
            dbg = bt_df_for_port.copy()

            st.write("columns:", list(dbg.columns))

            if "decision" in dbg.columns:
                st.write("decision counts")
                st.write(dbg["decision"].astype(str).value_counts(dropna=False))

            if "signal_action" in dbg.columns:
                st.write("signal_action counts")
                st.write(dbg["signal_action"].astype(str).value_counts(dropna=False))

            if "position_size_pct" in dbg.columns:
                st.write("position_size_pct counts")
                st.write(dbg["position_size_pct"].value_counts(dropna=False).sort_index())

            cols = [c for c in
                    ["as_of_date", "decision", "signal_action", "signal_confidence", "position_size_pct", "reason"] if
                    c in dbg.columns]
            st.write("head")
            st.dataframe(dbg[cols].head(15), use_container_width=True, hide_index=True)

        with st.expander("Primary bt_df_for_port debug", expanded=True):
            cols = [c for c in [
                "as_of_date",
                "decision",
                "signal_action",
                "position_size_pct",
                "signal_confidence",
                "reason",
            ] if c in bt_df_for_port.columns]
            st.dataframe(bt_df_for_port[cols].head(25), use_container_width=True, hide_index=True)

            if "signal_action" in bt_df_for_port.columns:
                st.write("signal_action counts")
                st.write(bt_df_for_port["signal_action"].astype(str).value_counts(dropna=False))

            if "position_size_pct" in bt_df_for_port.columns:
                st.write("position_size_pct counts")
                st.write(bt_df_for_port["position_size_pct"].value_counts(dropna=False).sort_index())

        portfolio_df = build_portfolio_backtest(
            bt_df_for_port,
            initial_capital=float(bt_initial_capital),
            trading_mode=bt_trading_mode,
            transaction_cost_bps=float(bt_tc_bps),
            slippage_bps=float(bt_slip_bps),
            execution_lag_steps=bt_execution_lag,
            min_hold_steps=bt_min_hold_steps,
            confirmation_steps=bt_confirmation_steps,
            rebalance_mode=bt_rebalance_mode,
            rebalance_threshold_pct=float(bt_rebalance_threshold),
            stop_loss_pct=float(bt_stop_loss_pct),
            take_profit_pct=float(bt_take_profit_pct),
        )

        with st.expander("DEBUG SUMMARY AFTER PORTFOLIO", expanded=True):
            dbg = portfolio_df.copy()

            st.write("columns:", list(dbg.columns))

            for col in ["raw_action", "confirmed_action", "executed_trade", "position_after_label"]:
                if col in dbg.columns:
                    st.write(f"{col} counts")
                    st.write(dbg[col].astype(str).value_counts(dropna=False))

            for col in ["position_after_pct", "applied_position_pct", "turnover_pct"]:
                if col in dbg.columns:
                    st.write(f"{col} counts")
                    st.write(dbg[col].value_counts(dropna=False).sort_index())

            cols = [c for c in [
                "as_of_date",
                "decision",
                "signal_action",
                "raw_action",
                "confirmed_action",
                "executed_trade",
                "position_before_pct",
                "position_after_pct",
                "applied_position_pct",
                "turnover_pct",
                "strategy_period_return_pct",
                "stop_triggered",
                "take_profit_triggered",
            ] if c in dbg.columns]

            st.write("head")
            st.dataframe(dbg[cols].head(20), use_container_width=True, hide_index=True)

        with st.expander("Primary portfolio_df debug", expanded=True):
            cols = [c for c in [
                "as_of_date",
                "decision",
                "signal_action",
                "position_size_pct",
                "raw_action",
                "confirmed_action",
                "signal_confirmed",
                "executed_trade",
                "position_before_label",
                "position_before_pct",
                "position_after_label",
                "position_after_pct",
                "applied_position_pct",
                "next_position_pct",
                "turnover_pct",
                "strategy_period_return_pct",
                "stop_triggered",
                "take_profit_triggered",
                "equity_curve",
            ] if c in portfolio_df.columns]
            st.dataframe(portfolio_df[cols].head(30), use_container_width=True, hide_index=True)

        bt_stats = compute_backtest_analytics(portfolio_df, periods_per_year=float(
            periods_per_year)) if not portfolio_df.empty else {}

        # ----------------------------
        # Run comparison backtest
        # ----------------------------
        compare_portfolio_df = None
        compare_stats = {}
        compare_bt_df = None

        if compare_enabled:
            compare_cfg = MacroBacktestConfig(
                ticker=bt_price_col,
                start_date=bt_start,
                step=bt_step,
                price_col=bt_price_col,
                forward_return_months=bt_fwd,
                trading_mode=compare_trading_mode,
                transaction_cost_bps=compare_tc_bps,
                slippage_bps=compare_slip_bps,
            )

            compare_bt_df = run_macro_strategy_backtest(
                res_base,
                compute_one_date=_compute_bt_row_factory(bt_compare_rule),
                config=compare_cfg,
            )

            # Force overwrite mapping so stale/empty columns cannot survive
            compare_bt_df["signal_action"] = compare_bt_df[
                "decision"] if "decision" in compare_bt_df.columns else np.nan
            compare_bt_df["signal_confidence"] = compare_bt_df[
                "decision_confidence"] if "decision_confidence" in compare_bt_df.columns else np.nan

            with st.expander("Comparison Replay Debug", expanded=False):
                st.write("Comparison bt_df columns:")
                st.write(list(compare_bt_df.columns))

                if "decision" in compare_bt_df.columns:
                    st.write("Comparison decision counts:")
                    st.write(compare_bt_df["decision"].astype(str).value_counts(dropna=False))
                    st.write("Comparison decision null count:", int(compare_bt_df["decision"].isna().sum()))
                else:
                    st.error("Comparison bt_df is missing the 'decision' column.")

                cmp_debug_cols = [c for c in [
                    "as_of_date",
                    "strategy_rule",
                    "gold_signal",
                    "gold_regime",
                    "market_signal",
                    "market_regime",
                    "accel_signal",
                    "accel_regime",
                    "macro_state",
                    "macro_score",
                    "decision",
                    "decision_confidence",
                    "position_size_pct",
                    "signal_action",
                    "signal_confidence",
                    "reason",
                    "error",
                ] if c in compare_bt_df.columns]

                if cmp_debug_cols:
                    st.write("Comparison replay sample:")
                    st.dataframe(compare_bt_df[cmp_debug_cols].head(25), width="stretch", hide_index=True)

                cmp_macro_debug_cols = [c for c in [
                    "as_of_date",
                    "gold_signal",
                    "gold_regime",
                    "market_signal",
                    "market_regime",
                    "accel_signal",
                    "accel_regime",
                    "macro_score",
                    "macro_state",
                    "decision",
                    "reason",
                ] if c in compare_bt_df.columns]

                if cmp_macro_debug_cols:
                    st.write("Comparison macro input sample:")
                    st.dataframe(compare_bt_df[cmp_macro_debug_cols].head(25), width="stretch", hide_index=True)

                if "error" in compare_bt_df.columns and compare_bt_df["error"].notna().any():
                    st.warning("Replay row errors detected in comparison strategy")
                    st.dataframe(
                        compare_bt_df.loc[compare_bt_df["error"].notna(), ["as_of_date", "error"]].head(20),
                        width="stretch",
                        hide_index=True,
                    )

            with st.expander("Comparison pre-portfolio debug", expanded=False):
                st.write("Comparison bt_df columns before _prepare_bt_for_portfolio:")
                st.write(list(compare_bt_df.columns))

            #---------- Comparison backtest start

            compare_bt_df_for_port = _prepare_bt_for_portfolio(compare_bt_df, compare_trading_mode)

            with st.expander("Comparison bt_df_for_port debug", expanded=True):
                dbg_cols = [c for c in [
                    "as_of_date",
                    "decision",
                    "signal_action",
                    "position_size_pct",
                    "signal_confidence",
                ] if c in compare_bt_df_for_port.columns]
                st.dataframe(compare_bt_df_for_port[dbg_cols].head(40), use_container_width=True)

            compare_portfolio_df = build_portfolio_backtest(
                compare_bt_df_for_port,
                initial_capital=float(bt_initial_capital),
                trading_mode=compare_trading_mode,
                transaction_cost_bps=float(compare_tc_bps),
                slippage_bps=float(compare_slip_bps),
                execution_lag_steps=bt_execution_lag,
                min_hold_steps=bt_min_hold_steps,
                confirmation_steps=bt_confirmation_steps,
                rebalance_mode=bt_rebalance_mode,
                rebalance_threshold_pct=float(bt_rebalance_threshold),
                stop_loss_pct=float(bt_stop_loss_pct),
                take_profit_pct=float(bt_take_profit_pct),
            )

            with st.expander("Comparison Portfolio Debug", expanded=True):
                dbg_cols = [c for c in [
                    "as_of_date",
                    "decision",
                    "signal_action",
                    "position_size_pct",
                    "raw_action",
                    "confirmed_action",
                    "signal_confirmed",
                    "executed_trade",
                    "position_before_pct",
                    "position_after_pct",
                    "applied_position_pct",
                    "next_position_pct",
                    "strategy_period_return_pct",
                    "equity_curve",
                ] if c in compare_portfolio_df.columns]

                st.dataframe(compare_portfolio_df[dbg_cols].head(40), use_container_width=True)

            # ---- DEBUG BLOCK (add here) ----
            with st.expander("Comparison Portfolio Debug", expanded=False):
                st.write("Comparison portfolio_df columns:")
                st.write(list(compare_portfolio_df.columns))

                dbg_cols = [c for c in [
                    "as_of_date",
                    "decision",
                    "signal_action",
                    "position_size_pct",
                    "executed_trade",
                    "position_before_pct",
                    "position_after_pct",
                    "position_before_label",
                    "position_after_label",
                    "strategy_period_return_pct",
                    "equity_curve",
                ] if c in compare_portfolio_df.columns]

                if dbg_cols:
                    st.dataframe(compare_portfolio_df[dbg_cols].head(25), width="stretch", hide_index=True)

            compare_stats = compute_backtest_analytics(
                compare_portfolio_df,
                periods_per_year=float(periods_per_year)
            ) if compare_portfolio_df is not None and not compare_portfolio_df.empty else {}

        # ----------------------------
        # Headline metrics
        # ----------------------------
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Replay rows", int(bt_stats.get("rows", 0)) if bt_stats else 0)
        m2.metric("BUY hit rate", f"{bt_stats.get('buy_hit_rate_pct', np.nan):.1f}%" if bt_stats and pd.notna(
            bt_stats.get("buy_hit_rate_pct", np.nan)) else "—")
        m3.metric("Strategy total return", f"{bt_stats.get('total_return_pct', np.nan):.1f}%" if bt_stats and pd.notna(
            bt_stats.get("total_return_pct", np.nan)) else "—")
        m4.metric("Max drawdown", f"{bt_stats.get('max_drawdown_pct', np.nan):.1f}%" if bt_stats and pd.notna(
            bt_stats.get("max_drawdown_pct", np.nan)) else "—")

        n1, n2, n3, n4 = st.columns(4)
        n1.metric("Confirmed rows", int(bt_stats.get("confirmed_rows", 0)) if bt_stats else 0)
        n2.metric("Stop events", int(bt_stats.get("stop_events", 0)) if bt_stats else 0)
        n3.metric("Take-profit events", int(bt_stats.get("take_profit_events", 0)) if bt_stats else 0)
        n4.metric("Avg turnover", f"{bt_stats.get('avg_turnover_pct', np.nan):.2f}%" if bt_stats and pd.notna(
            bt_stats.get("avg_turnover_pct", np.nan)) else "—")

        # ----------------------------
        # Comparison summary
        # ----------------------------
        if compare_enabled:
            st.markdown("### Side-by-Side Strategy Comparison")
            with st.expander("What this section means"):
                st.markdown("""
    This compares the two selected setups on the same historical period.

    You can compare:
    - different **rules**
    - different **trading modes**
    - different **cost/slippage assumptions**
                """)

            comp_rows = [
                {
                    "Setup": f"Primary — {bt_primary_rule}",
                    "Trading Mode": bt_trading_mode,
                    "TC (bps)": bt_tc_bps,
                    "Slip (bps)": bt_slip_bps,
                    "Total Return (%)": bt_stats.get("total_return_pct", np.nan),
                    "CAGR (%)": bt_stats.get("cagr_pct", np.nan),
                    "Volatility (%)": bt_stats.get("volatility_pct", np.nan),
                    "Sharpe-like": bt_stats.get("sharpe_like", np.nan),
                    "Max DD (%)": bt_stats.get("max_drawdown_pct", np.nan),
                },
                {
                    "Setup": f"Comparison — {bt_compare_rule}",
                    "Trading Mode": compare_trading_mode,
                    "TC (bps)": compare_tc_bps,
                    "Slip (bps)": compare_slip_bps,
                    "Total Return (%)": compare_stats.get("total_return_pct", np.nan),
                    "CAGR (%)": compare_stats.get("cagr_pct", np.nan),
                    "Volatility (%)": compare_stats.get("volatility_pct", np.nan),
                    "Sharpe-like": compare_stats.get("sharpe_like", np.nan),
                    "Max DD (%)": compare_stats.get("max_drawdown_pct", np.nan),
                },
            ]
            st.dataframe(pd.DataFrame(comp_rows), use_container_width=True, hide_index=True)

        # ----------------------------
        # Charts
        # ----------------------------
        if not portfolio_df.empty:
            curve_df = portfolio_df[["as_of_date", "equity_curve", "benchmark_curve", "price"]].dropna(how="all").copy()
            curve_df = curve_df.rename(columns={
                "equity_curve": "Primary Strategy",
                "benchmark_curve": "Benchmark Portfolio",
                "price": "Benchmark Raw Price",
            })

            if compare_enabled and compare_portfolio_df is not None and not compare_portfolio_df.empty:
                comp_curve = compare_portfolio_df[["as_of_date", "equity_curve"]].dropna().copy()
                comp_curve = comp_curve.rename(columns={"equity_curve": "Comparison Strategy"})
                curve_df = curve_df.merge(comp_curve, on="as_of_date", how="outer")

            curve_df = curve_df.set_index("as_of_date").sort_index()

            st.markdown("### Strategy vs Benchmark")
            with st.expander("What this chart means"):
                st.markdown("""
            This chart uses **3 axes**:

            - **Left axis** = strategy portfolio value
            - **Second left axis** = raw benchmark price
            - **Right axis** = cumulative % change since the start

            So you can read:
            - the actual money value of the strategies
            - the raw benchmark asset price
            - the percentage growth of each series

            The benchmark price is shown in its own raw market units, while the strategy values are shown in portfolio money terms.
                """)

            fig_curve = go.Figure()


            def _pct_from_start(s):
                s = pd.to_numeric(s, errors="coerce")
                first_valid = s.dropna()
                if first_valid.empty:
                    return s * np.nan
                start_val = float(first_valid.iloc[0])
                if start_val == 0:
                    return s * np.nan
                return (s / start_val - 1.0) * 100.0


            # ---------------------------
            # LEFT AXIS 1: strategy values
            # ---------------------------
            if "Primary Strategy" in curve_df.columns:
                fig_curve.add_trace(
                    go.Scatter(
                        x=curve_df.index,
                        y=curve_df["Primary Strategy"],
                        mode="lines",
                        name="Primary value",
                        yaxis="y",
                        line=dict(color="#1f77b4", width=2.5),
                        hovertemplate="%{x}<br>Primary value: %{y:,.2f}<extra></extra>",
                    )
                )

            if "Comparison Strategy" in curve_df.columns:
                fig_curve.add_trace(
                    go.Scatter(
                        x=curve_df.index,
                        y=curve_df["Comparison Strategy"],
                        mode="lines",
                        name="Comparison value",
                        yaxis="y",
                        line=dict(color="#66b3ff", width=2.5),
                        hovertemplate="%{x}<br>Comparison value: %{y:,.2f}<extra></extra>",
                    )
                )

            # ---------------------------
            # LEFT AXIS 2: raw benchmark price
            # ---------------------------
            if "Benchmark Raw Price" in curve_df.columns:
                fig_curve.add_trace(
                    go.Scatter(
                        x=curve_df.index,
                        y=curve_df["Benchmark Raw Price"],
                        mode="lines",
                        name="Benchmark raw price",
                        yaxis="y3",
                        line=dict(color="#d62728", width=2.5),
                        hovertemplate="%{x}<br>Benchmark raw price: %{y:,.2f}<extra></extra>",
                    )
                )

            # ---------------------------
            # RIGHT AXIS: % change
            # ---------------------------
            if "Primary Strategy" in curve_df.columns:
                prim_pct = _pct_from_start(curve_df["Primary Strategy"])
                fig_curve.add_trace(
                    go.Scatter(
                        x=curve_df.index,
                        y=prim_pct,
                        mode="lines",
                        name="Primary %",
                        yaxis="y2",
                        line=dict(color="#1f77b4", width=2, dash="dash"),
                        hovertemplate="%{x}<br>Primary change: %{y:.2f}%<extra></extra>",
                    )
                )

            if "Comparison Strategy" in curve_df.columns:
                comp_pct = _pct_from_start(curve_df["Comparison Strategy"])
                fig_curve.add_trace(
                    go.Scatter(
                        x=curve_df.index,
                        y=comp_pct,
                        mode="lines",
                        name="Comparison %",
                        yaxis="y2",
                        line=dict(color="#66b3ff", width=2, dash="dash"),
                        hovertemplate="%{x}<br>Comparison change: %{y:.2f}%<extra></extra>",
                    )
                )

            if "Benchmark Raw Price" in curve_df.columns:
                bench_pct = _pct_from_start(curve_df["Benchmark Raw Price"])
                fig_curve.add_trace(
                    go.Scatter(
                        x=curve_df.index,
                        y=bench_pct,
                        mode="lines",
                        name="Benchmark %",
                        yaxis="y2",
                        line=dict(color="#d62728", width=2, dash="dash"),
                        hovertemplate="%{x}<br>Benchmark change: %{y:.2f}%<extra></extra>",
                    )
                )

            fig_curve.update_layout(
                height=540,
                margin=dict(l=20, r=20, t=10, b=20),
                legend=dict(orientation="h", yanchor="bottom", y=-0.28, xanchor="left", x=0),
                hovermode="x unified",

                yaxis=dict(
                    title=dict(
                        text="Strategy portfolio value",
                        font=dict(color="#1f77b4"),
                    ),
                    tickfont=dict(color="#1f77b4"),
                    side="left",
                    showgrid=True,
                    zeroline=False,
                    showline=True,
                    linecolor="#1f77b4",
                    tickcolor="#1f77b4",
                ),

                yaxis2=dict(
                    title=dict(
                        text="% change since start",
                        font=dict(color="#666666"),
                    ),
                    tickfont=dict(color="#666666"),
                    overlaying="y",
                    side="right",
                    showgrid=False,
                    zeroline=False,
                    ticksuffix="%",
                    showline=True,
                    linecolor="#666666",
                    tickcolor="#666666",
                ),

                yaxis3=dict(
                    title=dict(
                        text="Benchmark raw price",
                        font=dict(color="#d62728"),
                    ),
                    tickfont=dict(color="#d62728"),
                    anchor="free",
                    overlaying="y",
                    side="left",
                    position=0.08,
                    showgrid=False,
                    zeroline=False,
                    showline=True,
                    linecolor="#d62728",
                    tickcolor="#d62728",
                ),
            )

            st.plotly_chart(fig_curve, use_container_width=True)

            dd_df = portfolio_df[["as_of_date", "drawdown_pct"]].dropna().copy().rename(
                columns={"drawdown_pct": "Primary Strategy"})
            if compare_enabled and compare_portfolio_df is not None and not compare_portfolio_df.empty:
                comp_dd = compare_portfolio_df[["as_of_date", "drawdown_pct"]].dropna().copy().rename(
                    columns={"drawdown_pct": "Comparison Strategy"})
                dd_df = dd_df.merge(comp_dd, on="as_of_date", how="outer")
            dd_df = dd_df.set_index("as_of_date").sort_index()

            st.markdown("### Strategy Drawdown (%)")
            with st.expander("What this chart means"):
                st.markdown("Smaller drops usually mean an easier ride emotionally and financially.")
            fig_dd = go.Figure()
            if "Comparison Strategy" in dd_df.columns:
                fig_dd.add_trace(
                    go.Scatter(x=dd_df.index, y=dd_df["Comparison Strategy"], mode="lines", name="Comparison Strategy",
                               line=dict(dash="dot")))
            if "Primary Strategy" in dd_df.columns:
                fig_dd.add_trace(
                    go.Scatter(x=dd_df.index, y=dd_df["Primary Strategy"], mode="lines", name="Primary Strategy",
                               line=dict(dash="dash")))
            fig_dd.update_layout(height=420, margin=dict(l=20, r=20, t=10, b=20), legend=dict(orientation="h"))
            st.plotly_chart(fig_dd, use_container_width=True)

            rets_df = portfolio_df[["as_of_date", "strategy_period_return_pct", "asset_period_return_pct"]].dropna(
                how="all").copy()
            rets_df = rets_df.rename(
                columns={"strategy_period_return_pct": "Primary Strategy", "asset_period_return_pct": "Asset"})
            if compare_enabled and compare_portfolio_df is not None and not compare_portfolio_df.empty:
                comp_rets = compare_portfolio_df[["as_of_date", "strategy_period_return_pct"]].dropna().copy()
                comp_rets = comp_rets.rename(columns={"strategy_period_return_pct": "Comparison Strategy"})
                rets_df = rets_df.merge(comp_rets, on="as_of_date", how="outer")
            rets_df = rets_df.set_index("as_of_date").sort_index()

            st.markdown("### Period Returns: Strategy vs Asset (%)")
            with st.expander("What this chart means"):
                st.markdown("""
    This shows the return earned during each replay step.

    - **Primary Strategy** = first setup
    - **Comparison Strategy** = second setup
    - **Asset** = the raw market move
                """)
            fig_rets = go.Figure()
            if "Asset" in rets_df.columns:
                fig_rets.add_trace(go.Scatter(x=rets_df.index, y=rets_df["Asset"], mode="lines", name="Asset"))
            if "Comparison Strategy" in rets_df.columns:
                fig_rets.add_trace(go.Scatter(x=rets_df.index, y=rets_df["Comparison Strategy"], mode="lines",
                                              name="Comparison Strategy", line=dict(dash="dot")))
            if "Primary Strategy" in rets_df.columns:
                fig_rets.add_trace(
                    go.Scatter(x=rets_df.index, y=rets_df["Primary Strategy"], mode="lines", name="Primary Strategy",
                               line=dict(dash="dash")))
            fig_rets.update_layout(height=420, margin=dict(l=20, r=20, t=10, b=20), legend=dict(orientation="h"))
            st.plotly_chart(fig_rets, use_container_width=True)

        # ----------------------------
        # Quick winner summary
        # ----------------------------
        st.markdown("### Quick Summary")
        with st.expander("What this summary means"):
            st.markdown("""
        This box highlights which setup looked best on three simple criteria:

        - **Highest total return** = made the most money overall
        - **Lowest max drawdown** = suffered the smallest worst drop
        - **Best return/drawdown tradeoff** = balanced gain versus pain most efficiently

        This helps you see quickly whether a strategy mainly wins on growth, safety, or balance.
            """)

        summary_candidates = [
            {
                "name": f"Primary — {bt_primary_rule}",
                "total_return_pct": bt_stats.get("total_return_pct", np.nan),
                "max_drawdown_pct": bt_stats.get("max_drawdown_pct", np.nan),
            }
        ]

        if compare_enabled and compare_stats:
            summary_candidates.append(
                {
                    "name": f"Comparison — {bt_compare_rule}",
                    "total_return_pct": compare_stats.get("total_return_pct", np.nan),
                    "max_drawdown_pct": compare_stats.get("max_drawdown_pct", np.nan),
                }
            )

        # Return / drawdown tradeoff:
        # use total_return / abs(max_drawdown), but only when drawdown is valid and non-zero
        for item in summary_candidates:
            tr = pd.to_numeric(pd.Series([item["total_return_pct"]]), errors="coerce").iloc[0]
            dd = pd.to_numeric(pd.Series([item["max_drawdown_pct"]]), errors="coerce").iloc[0]
            if pd.notna(tr) and pd.notna(dd) and abs(dd) > 1e-9:
                item["tradeoff_score"] = float(tr) / abs(float(dd))
            else:
                item["tradeoff_score"] = np.nan


        def _best_valid(items, key, higher_is_better=True):
            vals = []
            for x in items:
                v = pd.to_numeric(pd.Series([x.get(key)]), errors="coerce").iloc[0]
                if pd.notna(v):
                    vals.append((x["name"], float(v)))
            if not vals:
                return ("—", np.nan)
            return max(vals, key=lambda z: z[1]) if higher_is_better else min(vals, key=lambda z: z[1])


        best_return_name, best_return_val = _best_valid(summary_candidates, "total_return_pct", higher_is_better=True)
        best_drawdown_name, best_drawdown_val = _best_valid(summary_candidates, "max_drawdown_pct",
                                                            higher_is_better=True)
        best_tradeoff_name, best_tradeoff_val = _best_valid(summary_candidates, "tradeoff_score", higher_is_better=True)

        s1, s2, s3 = st.columns(3)
        with s1:
            st.info(
                f"**Highest total return**\n\n"
                f"{best_return_name}\n\n"
                f"{best_return_val:.1f}%" if pd.notna(best_return_val) else
                "**Highest total return**\n\n—"
            )

        with s2:
            st.info(
                f"**Lowest max drawdown**\n\n"
                f"{best_drawdown_name}\n\n"
                f"{best_drawdown_val:.1f}%" if pd.notna(best_drawdown_val) else
                "**Lowest max drawdown**\n\n—"
            )

        with s3:
            st.info(
                f"**Best return / drawdown tradeoff**\n\n"
                f"{best_tradeoff_name}\n\n"
                f"{best_tradeoff_val:.2f}" if pd.notna(best_tradeoff_val) else
                "**Best return / drawdown tradeoff**\n\n—"
            )

        # ----------------------------
        # Historical Monte Carlo Validation results
        # ----------------------------
        if wf_mc_enabled and mc_validation_summary:
            st.markdown("### Historical Monte Carlo Validation Results")
            with st.expander("What these metrics mean"):
                st.markdown("""
        These metrics tell you how well the Monte Carlo cone worked historically.

        - **Inside p10–p90** = how often the real later outcome stayed inside the forecast band
        - **Below p10** = how often reality was worse than the model's lower band
        - **Above p90** = how often reality was stronger than the model's upper band
        - **Directional hit** = how often the sign of the actual return matched the forecast median
        - **Median abs error vs p50** = typical distance between reality and the model's middle forecast
                """)

            vm1, vm2, vm3, vm4, vm5 = st.columns(5)
            vm1.metric("Inside p10–p90", f"{mc_validation_summary.get('inside_p10_p90_pct', np.nan):.1f}%")
            vm2.metric("Below p10", f"{mc_validation_summary.get('below_p10_pct', np.nan):.1f}%")
            vm3.metric("Above p90", f"{mc_validation_summary.get('above_p90_pct', np.nan):.1f}%")
            vm4.metric("Directional hit", f"{mc_validation_summary.get('direction_hit_pct', np.nan):.1f}%")
            vm5.metric("Median abs error vs p50", f"{mc_validation_summary.get('median_abs_err_pct', np.nan):.2f}%")

        if wf_mc_enabled and not bt_df.empty and {"mc_p10_ret_pct", "mc_p50_ret_pct", "mc_p90_ret_pct",
                                                  "mc_actual_ret_pct"}.issubset(bt_df.columns):

            col_mc_left, col_mc_right = st.columns(2)

            # -------------------------------------------------
            # LEFT: Existing historical validation chart
            # -------------------------------------------------
            with col_mc_left:
                val_plot = bt_df[
                    ["as_of_date", "mc_p10_ret_pct", "mc_p50_ret_pct", "mc_p90_ret_pct", "mc_actual_ret_pct"]
                ].dropna(how="all").copy()

                if not val_plot.empty:
                    val_plot["as_of_date"] = pd.to_datetime(val_plot["as_of_date"], errors="coerce")
                    val_plot = val_plot.sort_values("as_of_date")

                    st.markdown("### Forecast Cone vs Actual Outcome")
                    with st.expander("What this chart means"):
                        st.markdown("""
        For each replay date:
        - dashed lower line = Monte Carlo p10
        - middle dashed line = Monte Carlo p50
        - dashed upper line = Monte Carlo p90
        - solid line = what actually happened later

        If the solid line often stays between p10 and p90, the forecast cone is reasonably calibrated.
                        """)

                    fig_val = go.Figure()

                    fig_val.add_trace(go.Scatter(
                        x=val_plot["as_of_date"],
                        y=val_plot["mc_p10_ret_pct"],
                        mode="lines",
                        name="MC p10",
                        line=dict(color="#d62728", dash="dash"),
                    ))
                    fig_val.add_trace(go.Scatter(
                        x=val_plot["as_of_date"],
                        y=val_plot["mc_p50_ret_pct"],
                        mode="lines",
                        name="MC p50",
                        line=dict(color="#1f77b4", dash="dash"),
                    ))
                    fig_val.add_trace(go.Scatter(
                        x=val_plot["as_of_date"],
                        y=val_plot["mc_p90_ret_pct"],
                        mode="lines",
                        name="MC p90",
                        line=dict(color="#2ca02c", dash="dash"),
                    ))
                    fig_val.add_trace(go.Scatter(
                        x=val_plot["as_of_date"],
                        y=val_plot["mc_actual_ret_pct"],
                        mode="lines",
                        name="Actual forward return",
                        line=dict(color="#000000", width=2),
                    ))

                    fig_val.update_layout(
                        height=420,
                        margin=dict(l=20, r=20, t=10, b=20),
                        legend=dict(orientation="h"),
                        yaxis=dict(title=f"Forward return over {bt_fwd} bars (%)"),
                        hovermode="x unified",
                    )
                    st.plotly_chart(fig_val, use_container_width=True)

            # -------------------------------------------------
            # RIGHT: Historical cone replay from Start date
            # -------------------------------------------------
            with col_mc_right:
                st.markdown("### Historical Cone Replay from Start Date")
                with st.expander("What this chart means"):
                    st.markdown("""
        This chart anchors the Monte Carlo cone at the Walk-Forward **Start date**.

        It shows:
        - recent actual price history leading into that date
        - the simulated Monte Carlo cone from that date forward
        - the actual realized price path over the next forecast horizon

        This lets you judge visually how accurate the cone was for that specific historical start date.
                    """)

                price_full = pd.to_numeric(res_base[bt_price_col], errors="coerce").dropna().astype(float).sort_index()
                anchor_ts = pd.Timestamp(bt_start)

                hist_px = price_full.loc[price_full.index <= anchor_ts].copy()

                if hist_px.empty:
                    st.info("No price history available up to the selected Start date.")
                else:
                    anchor_used = hist_px.index[-1]
                    start_price = float(hist_px.iloc[-1])

                    # recent history for visual context
                    hist_context = price_full.loc[price_full.index <= anchor_used].tail(12).copy()

                    # realized future path over the next bt_fwd bars
                    actual_future = price_full.loc[price_full.index > anchor_used].head(int(bt_fwd)).copy()

                    # benchmark history for tactical-state construction
                    bench_hist = None
                    if "QQQ" in res_base.columns:
                        bench_full = pd.to_numeric(res_base["QQQ"], errors="coerce").dropna().astype(float).sort_index()
                        bench_hist = bench_full.loc[bench_full.index <= anchor_used].copy()

                    cone_cfg = MCValidationConfig(
                        horizon_steps=int(bt_fwd),
                        n_sims=int(wf_mc_sims),
                        block_size=int(wf_mc_block),
                        seed=7,
                    )

                    state_hist, meta = build_tactical_state_series(
                        price=hist_px,
                        benchmark_price=bench_hist,
                        cfg=cone_cfg,
                    )

                    cone_df, state_used = mc.monte_carlo_cone_by_tactical_state_block(
                        price=hist_px,
                        state=state_hist,
                        state_now=meta.get("state_now"),
                        horizon_steps=int(bt_fwd),
                        n_sims=int(wf_mc_sims),
                        block_size=int(wf_mc_block),
                        seed=7,
                    )

                    if cone_df.empty:
                        st.info(f"Not enough history in tactical state ({state_used}) to build the cone.")
                    else:
                        # Use actual future dates when available; otherwise fallback to monthly future dates
                        if len(actual_future) >= len(cone_df):
                            future_dates = actual_future.index[:len(cone_df)]
                        else:
                            future_dates = pd.date_range(
                                start=anchor_used,
                                periods=len(cone_df) + 1,
                                freq="M"
                            )[1:]

                        cone_plot = cone_df.copy()
                        cone_plot["future_date"] = future_dates

                        # prepend anchor point so p10/p50/p90 all start from same price
                        anchor_row = pd.DataFrame({
                            "future_date": [anchor_used],
                            "p10": [start_price],
                            "p50": [start_price],
                            "p90": [start_price],
                        })
                        cone_plot = pd.concat(
                            [anchor_row, cone_plot[["future_date", "p10", "p50", "p90"]]],
                            ignore_index=True
                        )

                        # actual realized path also starts from the anchor
                        actual_plot = pd.concat([
                            pd.Series([start_price], index=[anchor_used]),
                            actual_future
                        ]).sort_index()

                        fig_cone = go.Figure()

                        # recent history before anchor
                        fig_cone.add_trace(go.Scatter(
                            x=hist_context.index,
                            y=hist_context.values,
                            mode="lines",
                            name="Actual price (pre-anchor)",
                            line=dict(color="#7f7f7f"),
                        ))

                        # upper band first
                        fig_cone.add_trace(go.Scatter(
                            x=cone_plot["future_date"],
                            y=cone_plot["p90"],
                            mode="lines",
                            name="MC p90",
                            line=dict(color="#2ca02c", dash="dash"),
                        ))

                        # lower band, filled to previous trace
                        fig_cone.add_trace(go.Scatter(
                            x=cone_plot["future_date"],
                            y=cone_plot["p10"],
                            mode="lines",
                            name="MC p10",
                            line=dict(color="#d62728", dash="dash"),
                            fill="tonexty",
                            fillcolor="rgba(31,119,180,0.10)",
                        ))

                        # median path
                        fig_cone.add_trace(go.Scatter(
                            x=cone_plot["future_date"],
                            y=cone_plot["p50"],
                            mode="lines",
                            name="MC p50",
                            line=dict(color="#1f77b4", dash="dash"),
                        ))

                        # actual realized path
                        fig_cone.add_trace(go.Scatter(
                            x=actual_plot.index,
                            y=actual_plot.values,
                            mode="lines+markers",
                            name="Actual realized path",
                            line=dict(color="#000000", width=2),
                        ))

                        fig_cone.update_layout(
                            height=420,
                            margin=dict(l=20, r=20, t=10, b=20),
                            legend=dict(orientation="h"),
                            yaxis=dict(title="Price"),
                            hovermode="x unified",
                        )

                        st.caption(f"Anchor used: {anchor_used.date()} | MC tactical state: {state_used}")
                        st.plotly_chart(fig_cone, use_container_width=True)

        if wf_mc_enabled and not bt_df.empty:
            mc_cols = [
                "as_of_date",
                "mc_state",
                "mc_p10_ret_pct",
                "mc_p50_ret_pct",
                "mc_p90_ret_pct",
                "mc_prob_up_pct",
                "mc_actual_ret_pct",
                "mc_inside_p10_p90",
                "mc_below_p10",
                "mc_above_p90",
                "mc_direction_hit",
                "mc_abs_err_p50",
            ]
            mc_cols = [c for c in mc_cols if c in bt_df.columns]

            if mc_cols:
                st.markdown("### Historical MC Validation Table")
                with st.expander("What this table means"):
                    st.markdown("""
        This table shows the Monte Carlo forecast and the actual later result for each replay date.

        It helps you inspect where the forecast worked well, where it missed, and whether errors happened in specific market states.
                    """)
                st.dataframe(bt_df[mc_cols], use_container_width=True, hide_index=True)

        # ----------------------------
        # Tables
        # ----------------------------
        if not portfolio_df.empty:
            st.markdown("### Action Performance Summary")
            with st.expander("What this table means"):
                st.markdown("""
    This table groups historical results by the action the rule suggested.

    It helps you check whether BUY, WATCH, HOLD, and SELL were actually useful on average.
                """)

            fwd_cols = [c for c in portfolio_df.columns if c.startswith("fwd_") and c.endswith("m_ret_pct")]
            fwd_col = fwd_cols[0] if fwd_cols else None
            summary_rows = []
            if fwd_col is not None:
                tmp = portfolio_df.copy()
                tmp["decision"] = tmp["decision"].astype(str)
                for act, grp in tmp.groupby("decision", dropna=False):
                    vals = pd.to_numeric(grp[fwd_col], errors="coerce").dropna()
                    summary_rows.append({
                        "Action": act,
                        "Rows": int(len(grp)),
                        "Avg forward return (%)": float(vals.mean()) if not vals.empty else np.nan,
                        "Median forward return (%)": float(vals.median()) if not vals.empty else np.nan,
                        "Hit rate (%)": float((vals > 0).mean() * 100.0) if not vals.empty else np.nan,
                    })
            st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

            st.markdown("### Replay Table")
            with st.expander("Signals vs executed positions — plain-English guide"):
                st.markdown("""
    There are **two layers** here:

    - **Signal Action** = what the rule suggested
    - **Executed Trade / Position** = what the portfolio actually did after applying lag, confirmation, holding, and risk rules

    So a BUY signal does not always mean the trade was executed immediately.
                """)

            show_cols = [
                "as_of_date",
                "strategy_rule",
                "gold_regime",
                "market_regime",
                "accel_regime",
                "macro_state",
                "macro_score",
                "decision",
                "decision_confidence",
                "signal_action",
                "executed_trade",
                "position_before_label",
                "position_before_pct",
                "position_after_label",
                "position_after_pct",
                "hold_steps_in_position",
                "stop_triggered",
                "take_profit_triggered",
                "asset_period_return_pct",
                "strategy_period_return_pct",
                "equity_curve",
                "benchmark_curve",
                "drawdown_pct",
                "reason",
            ]
            show_cols = [c for c in show_cols if c in portfolio_df.columns]
            replay_show = portfolio_df[show_cols].copy()
            st.dataframe(replay_show, use_container_width=True, hide_index=True)

            if compare_enabled and compare_bt_df is not None and not compare_bt_df.empty:
                cmp_macro_debug_cols = [c for c in [
                    "as_of_date",
                    "gold_signal",
                    "gold_regime",
                    "market_signal",
                    "market_regime",
                    "accel_signal",
                    "accel_regime",
                    "macro_score",
                    "macro_state",
                    "decision",
                    "reason",
                ] if c in compare_bt_df.columns]

                st.write("Comparison macro input sample:")
                st.dataframe(
                    compare_bt_df[cmp_macro_debug_cols].head(25),
                    width="stretch",
                    hide_index=True,
                )

    # ---- Gold-only panels must render ONLY in Tab1
    with tab1:
        bull_now, bear_now = render_gold_action_panels(
            res=res,
            res_base=res_base,
            df=df,
            mode=mode,
            win_start=None,
            win_end=None,
            labels=labels,
            thresholds=thresholds,
            core_keys=core_keys,
            latest=latest,
            lookback=lookback,
            show_indicator_thresholds=show_indicator_thresholds,
            history_view=history_view,
            trig_hi=trig_hi,
            trig_lo=trig_lo,
            persist=persist,
            latest_market=latest_market,
            core_keys_market=core_keys_market,
            thresholds_market=thresholds_market,
            res_market=res_market,
        )

st.markdown("---")
st.caption("v4 notes: Real yield + inflation expectations use long-history proxies to support 1970s/1980s presets. You can later swap to TIPS real yields and forward inflation expectations when you want modern purity.")

st.markdown("## Report")

legend_map = {
"REAL_YIELD_CPI": "Real yield (10Y − CPI YoY)",
"CPI_YOY": "Inflation (CPI YoY)",
"USD_12M_CHG": "USD 12M % change (TWEX)",
"CURVE_10Y_3M": "Curve (10Y–3M)",
"DEFICIT_GDP": "Deficit % GDP",
"REAL_YIELD_TIPS10": "Real yield (10Y TIPS, 20D MA)",
"HY_OAS": "High Yield OAS (20D MA)",
}

# --- Safety defaults: ensure these always exist (prevents NameError in exports)
bull_now = globals().get("bull_now", False)
bear_now = globals().get("bear_now", False)
trig_hi  = globals().get("trig_hi", 0.60)
trig_lo  = globals().get("trig_lo", -0.60)
persist  = globals().get("persist", 2)

trigger_info = {
"bull_now": bull_now,
"bear_now": bear_now,
"trig_hi": trig_hi,
"trig_lo": trig_lo,
"persist": persist,
}

if st.button("Generate Word report (.docx)"):
    if enable_compare and compare_modes and len(compare_modes) == 2:
        left_mode, right_mode = compare_modes[0], compare_modes[1]

        def compute_mode_pack(m: str):
            if m == "Structural Regime (today)":
                return run_structural(res_base, weights_struct, dirs_struct), labels

            if m == "Crisis Similarity (template)":
                if not (win_start and win_end):
                    st.error("Comparison report includes Crisis Similarity but no window selected.")
                    st.stop()
                pack6 = run_crisis(res_base, weights_crisis, dirs_crisis, win_start, win_end)
                pack5 = pack6[:5]  # drop gold_stats for compare report
                return pack5, labels

            # Market Acceleration
            return run_accel(res_accel_base, accel_method, w_g, w_ry, w_st, w_u), labels_accel

        packL, labelsL = compute_mode_pack(left_mode)
        packR, labelsR = compute_mode_pack(right_mode)

        docx_bytes = build_word_report_compare(
            left_title=left_mode,
            right_title=right_mode,
            left_pack=packL,
            right_pack=packR,
            trigger_info=trigger_info,
            labels_left=labelsL,
            labels_right=labelsR,
            legend_map=legend_map,
        )

        fname = f"Gold_Macro_Cockpit_Compare_{left_mode[:10]}_{right_mode[:10]}_{latest.name.date()}.docx"

    else:
        # single-mode report
        end_dt = res.index.max()
        if history_view == "Full history":
            res_plot_for_report = res
        elif history_view == "Last 15y":
            res_plot_for_report = res.loc[end_dt - pd.DateOffset(years=15):]
        elif history_view == "Last 5y":
            res_plot_for_report = res.loc[end_dt - pd.DateOffset(years=5):]
        elif history_view == "Crisis window only" and (win_start and win_end):
            res_plot_for_report = res.loc[pd.to_datetime(win_start):pd.to_datetime(win_end)]
        elif history_view == "Crisis window ±5y" and (win_start and win_end):
            start_dt = pd.to_datetime(win_start) - pd.DateOffset(years=5)
            end_win = pd.to_datetime(win_end) + pd.DateOffset(years=5)
            res_plot_for_report = res.loc[start_dt:end_win]
        else:
            res_plot_for_report = res

        docx_bytes = build_word_report(
            latest_row=latest,
            prev_row=prev,
            thresholds=thresholds,
            weights=weights,
            trigger_info=trigger_info,
            labels=labels,
            legend_map=legend_map,
            core_keys=core_keys,
            mode=mode,
            crisis_year=crisis,
            res_plot=res_plot_for_report,
            gold_stats=gold_stats if mode == "Crisis Similarity (template)" else None
        )
        fname = f"Gold_Macro_Cockpit_Report_{latest.name.date()}.docx"

    st.download_button(
        label="Download report",
        data=docx_bytes,
        file_name=fname,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
